"""
Gera documentacao completa final do CongregaFiel
Sprints 1 a 6 - Documento unificado com capa, folha de rosto, sumario
Baseado no PDF 'CongregaFiel - Documentacao Completa.pdf' + Sprint 6
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ==================== ESTILOS GLOBAIS ====================
FONTE = 'Times New Roman'

style = doc.styles['Normal']
style.font.name = FONTE
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = FONTE
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)

# Margens ABNT
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)


# ==================== FUNCOES AUXILIARES ====================
def cover_line(text, size=12, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = FONTE
    return p


def h1(text):
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = FONTE
        run.bold = True
    return h


def h2(text):
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = FONTE
        run.bold = True
    return h


def body(text, bold=False, indent=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = FONTE
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.clear()
    run = p.add_run(text)
    run.font.name = FONTE
    run.font.size = Pt(12)
    return p


def tabela(headers, data):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = FONTE
                run.font.size = Pt(10)
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = FONTE
                    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return table


def spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def add_toc():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld1)

    run2 = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._r.append(instr)

    run3 = p.add_run()
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fld2)

    run4 = p.add_run('(Clique com botao direito e selecione "Atualizar campo", ou pressione F9)')
    run4.font.name = FONTE
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(128, 128, 128)

    run5 = p.add_run()
    fld3 = OxmlElement('w:fldChar')
    fld3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fld3)


# ==================== CAPA ====================
spacer(5)
cover_line('FACULDADE INSTED', 14, True, 4)
cover_line('Curso Superior de Tecnologia em Analise e Desenvolvimento de Sistemas', 12, False, 0)
spacer(6)
cover_line('CONGREGA FIEL', 16, True, 4)
cover_line('Sistema Web para Gestao de Comunidades Eclesiasticas', 13, False, 0)
spacer(4)
cover_line('Documentacao Tecnica Completa', 13, True, 0)
spacer(6)
for nome in ['Catieli Gama Cora', 'Fernando Alves da Nobrega',
             'Gabriel Franklin Barcellos', 'Jhenniffer Lopes da Silva Vargas',
             'Joao Pedro Aranda']:
    cover_line(nome, 12, False, 2)
spacer(4)
cover_line('Campo Grande - MS', 12, False, 2)
cover_line('2026', 12, False, 0)
doc.add_page_break()

# ==================== FOLHA DE ROSTO ====================
spacer(5)
cover_line('CONGREGA FIEL', 16, True, 4)
cover_line('Sistema Web para Gestao de Comunidades Eclesiasticas', 13, False, 0)
spacer(4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(0)
run = p.add_run(
    'Documentacao tecnica completa apresentada como requisito parcial para aprovacao '
    'no Curso Superior de Tecnologia em Analise e Desenvolvimento de Sistemas da '
    'FACULDADE INSTED.'
)
run.font.name = FONTE
run.font.size = Pt(11)
spacer(6)
cover_line('Equipe de Desenvolvimento', 12, True, 6)
for nome in ['Catieli Gama Cora', 'Fernando Alves da Nobrega',
             'Gabriel Franklin Barcellos', 'Jhenniffer Lopes da Silva Vargas',
             'Joao Pedro Aranda']:
    cover_line(nome, 12, False, 2)
spacer(4)
cover_line('Campo Grande - MS', 12, False, 2)
cover_line('2026', 12, False, 0)
doc.add_page_break()

# ==================== SUMARIO ====================
h1('SUMARIO')
add_toc()
doc.add_page_break()

# ============================================================
# 1 INTRODUCAO
# ============================================================
h1('1  INTRODUCAO')
body(
    'O Congrega Fiel e um sistema web para gestao de comunidades eclesiasticas, desenvolvido como '
    'projeto integrador do curso de Analise e Desenvolvimento de Sistemas da Faculdade INSTED. O sistema '
    'funciona como uma plataforma privada onde pastores gerenciam membros, eventos, contribuicoes e '
    'comunicados de suas igrejas.',
    indent=True
)
body(
    'O desenvolvimento foi organizado em seis sprints semanais, de 24 de fevereiro a 06 de abril de '
    '2026. Este documento consolida as entregas de todas as sprints.',
    indent=True
)

h2('1.1  Stack Tecnologica')
bullet('Frontend: HTML5, CSS3 e JavaScript puro (sem frameworks)')
bullet('Backend: Express.js (Node.js) e FastAPI (Python)')
bullet('Banco de dados: Supabase (PostgreSQL na nuvem)')
bullet('Hospedagem: Firebase Hosting (frontend), Vercel (APIs)')
bullet('Testes: Vitest')
bullet('Infraestrutura: API Gateway Edge, Microservicos')

h2('1.2  Equipe')
tabela(
    ['Integrante', 'Area'],
    [
        ['Catieli Gama Cora', 'Documentacao, Front-End'],
        ['Fernando Alves da Nobrega', 'Documentacao, Front-End, Capacitacao'],
        ['Gabriel Franklin Barcellos', 'Front-End, Back-End, API Gateway'],
        ['Jhenniffer Lopes da Silva Vargas', 'Documentacao, Front-End'],
        ['Joao Pedro Aranda', 'Back-End, Testes, API Gateway'],
    ]
)

# ============================================================
# 2 SPRINT 1
# ============================================================
doc.add_page_break()
h1('2  SPRINT 1 - DOCUMENTACAO DO MVP')
body('Periodo: 24 de fevereiro a 02 de marco de 2026.', italic=True)
body(
    'A primeira sprint definiu o escopo do Produto Minimo Viavel, os requisitos do sistema e a '
    'arquitetura tecnica.',
    indent=True
)

h2('2.1  Problema Identificado')
body(
    'Igrejas de pequeno e medio porte enfrentam dificuldades na gestao: cadastro manual de membros, '
    'controle financeiro precario, comunicacao fragmentada via redes sociais e ausencia de uma plataforma '
    'centralizada.',
    indent=True
)

h2('2.2  Solucao Proposta')
body('Plataforma web privada com dois perfis de acesso:', indent=True)
bullet('Pastor (administrador): cadastra igreja, gerencia membros, registra pagamentos, cria eventos e comunicados')
bullet('Fiel (membro): acessa perfil, eventos, contribuicoes e comunicados da igreja')

h2('2.3  Modulos do MVP')
tabela(
    ['Modulo', 'Descricao'],
    [
        ['Autenticacao', 'Cadastro, login, recuperacao de senha, perfis diferenciados'],
        ['Gestao da Igreja', 'Cadastro e configuracao da igreja pelo pastor'],
        ['Gestao de Membros', 'Cadastro, listagem, edicao e exclusao de membros'],
        ['Financeiro', 'Registro de contribuicoes (dizimos, ofertas, doacoes)'],
        ['Eventos', 'Criacao e listagem de eventos da comunidade'],
        ['Comunicados', 'Publicacao de avisos para os membros'],
    ]
)

h2('2.4  Arquitetura')
body(
    'O sistema adota arquitetura REST (Representational State Transfer) com comunicacao via HTTP '
    'e dados em formato JSON. Na fase inicial, utilizou-se JSON Server para simular a API e LocalStorage '
    'para persistencia local.',
    indent=True
)

# ============================================================
# 3 SPRINT 2
# ============================================================
doc.add_page_break()
h1('3  SPRINT 2 - DEPLOY ONLINE E BANCO DE DADOS')
body('Periodo: 03 a 09 de marco de 2026.', italic=True)
body(
    'A Sprint 2 migrou o sistema do ambiente local para a internet, com banco de dados real e duas '
    'APIs REST.',
    indent=True
)

h2('3.1  Entregas')
bullet('Site publicado no Firebase Hosting (HTTPS automatico)')
bullet('Banco de dados PostgreSQL no Supabase com 6 tabelas e Row Level Security')
bullet('API REST com Express.js (33 rotas) deployada no Vercel')
bullet('API REST com FastAPI (33 rotas) deployada no Vercel, com Swagger automatico')
bullet('Sistema de autenticacao com Supabase Auth (JWT, criptografia bcrypt)')
bullet('Servicos compartilhados no frontend (sessao, API, interface)')

h2('3.2  Banco de Dados')
tabela(
    ['Tabela', 'Finalidade'],
    [
        ['igrejas', 'Dados das igrejas cadastradas'],
        ['membros', 'Fieis vinculados a cada igreja'],
        ['eventos', 'Programacao de atividades'],
        ['contribuicoes', 'Registros financeiros (dizimos, ofertas)'],
        ['comunicados', 'Avisos e informes da igreja'],
        ['pedidos_oracao', 'Pedidos de oracao dos membros'],
    ]
)

h2('3.3  Comparativo dos Frameworks')
tabela(
    ['Criterio', 'Express.js', 'FastAPI'],
    [
        ['Linguagem', 'JavaScript', 'Python'],
        ['Validacao', 'Manual', 'Automatica (Pydantic)'],
        ['Documentacao', 'Manual', 'Swagger automatico'],
        ['Performance', 'Boa', 'Superior (async nativo)'],
    ]
)

h2('3.4  URLs de Producao')
tabela(
    ['Componente', 'URL'],
    [
        ['Frontend', 'https://congregafiel.web.app'],
        ['API Express', 'https://api-express-tau.vercel.app'],
        ['API FastAPI', 'https://api-fastapi.vercel.app'],
        ['Swagger', 'https://api-fastapi.vercel.app/docs'],
    ]
)

# ============================================================
# 4 SPRINT 3
# ============================================================
doc.add_page_break()
h1('4  SPRINT 3 - WEB SERVICES E MELHORIAS')
body('Periodo: 10 a 16 de marco de 2026.', italic=True)
body(
    'A Sprint 3 integrou servicos web externos e melhorou a experiencia do usuario com recursos '
    'visuais interativos.',
    indent=True
)

h2('4.1  Mapa Interativo (Web Service)')
body(
    'Integracao com OpenStreetMap via Leaflet.js na pagina de cadastro de membros. O mapa exibe '
    'marcadores das igrejas cadastradas, permite busca por nome e utiliza a Geolocation API do navegador '
    'para ordenar igrejas por proximidade (formula de Haversine).',
    indent=True
)

h2('4.2  QR Code para Convite')
body(
    'Geracao de QR Code no painel da igreja usando a biblioteca qrcode-generator (CDN). O codigo '
    'codifica um link direto para cadastro com a igreja pre-selecionada. Opcoes: exibir QR Code, copiar link '
    'ou baixar imagem PNG.',
    indent=True
)

h2('4.3  Linha do Tempo de Eventos')
body(
    'Nova pagina com visualizacao cronologica vertical dos eventos, organizada por mes. Inclui '
    'contagem regressiva para o proximo evento, filtros por tipo e animacoes de entrada com Intersection '
    'Observer.',
    indent=True
)

h2('4.4  Categorizacao de Eventos')
tabela(
    ['Tipo', 'Cor', 'Exemplo'],
    [
        ['Culto', 'Marrom (#D4A574)', 'Culto de Domingo'],
        ['Estudo Biblico', 'Azul (#5B8DEF)', 'Estudo de Romanos'],
        ['Conferencia', 'Roxo (#9B6FD9)', 'Conferencia de Jovens'],
        ['Evento Especial', 'Rosa (#E87C8A)', 'Aniversario da Igreja'],
        ['Outro', 'Bege (#C8956C)', 'Mutirao de Limpeza'],
    ]
)

h2('4.5  Outras Melhorias')
bullet('Carrossel horizontal de eventos nos paineis (igreja e membro)')
bullet('Menu lateral atualizado com "Linha do Tempo" e "Meu Perfil"')
bullet('Pagina de perfil do membro com edicao de dados e troca de senha')
bullet('Filtros por status nos pedidos de oracao (Todos, Pendentes, Orados, Respondidos)')
bullet('Paginas de comunicados e pagamentos aprimoradas')

h2('4.6  Web Services Utilizados')
tabela(
    ['Tecnologia', 'Tipo', 'Finalidade'],
    [
        ['OpenStreetMap', 'Web Service REST', 'Tiles de mapa'],
        ['Leaflet.js', 'Biblioteca JS (CDN)', 'Renderizacao do mapa'],
        ['Geolocation API', 'API do navegador', 'Posicao do usuario'],
        ['qrcode-generator', 'Biblioteca JS (CDN)', 'Geracao de QR Code'],
        ['Intersection Observer', 'API do navegador', 'Animacoes de scroll'],
    ]
)

# ============================================================
# 5 SPRINT 4
# ============================================================
doc.add_page_break()
h1('5  SPRINT 4 - CONSOLIDACAO E REFINAMENTO')
body('Periodo: 17 a 23 de marco de 2026.', italic=True)
body(
    'A Sprint 4 consolidou e refinou as funcionalidades da Sprint 3, com foco em testes, ajustes de '
    'responsividade e validacao dos fluxos completos.',
    indent=True
)

h2('5.1  Refinamentos')
bullet('Responsividade do mapa interativo para dispositivos moveis')
bullet('Otimizacao das animacoes da linha do tempo')
bullet('Validacao do fluxo completo de convite via QR Code')
bullet('Sistema de respostas nos pedidos de oracao')
bullet('Ajustes na pagina de perfil do membro')

# ============================================================
# 6 SPRINT 5
# ============================================================
doc.add_page_break()
h1('6  SPRINT 5 - TESTES UNITARIOS')
body('Periodo: 24 a 30 de marco de 2026.', italic=True)
body(
    'A Sprint 5 implementou testes unitarios com o framework Vitest, cobrindo os servicos centrais do '
    'sistema.',
    indent=True
)

h2('6.1  Framework')
body(
    'O Vitest e um framework de testes compativel com a sintaxe do Jest, otimizado para projetos '
    'JavaScript. Execucao via comando: npm test.',
    indent=True
)

h2('6.2  Testes Implementados')
tabela(
    ['Camada', 'Arquivo de Teste', 'Cobertura'],
    [
        ['Backend', 'pedidos-oracao-utils.test.js', 'Utilidades de pedidos de oracao'],
        ['Backend', 'regras-auth.test.js', 'Regras de autenticacao'],
        ['Frontend', 'api-servico.test.js', 'Comunicacao com API'],
        ['Frontend', 'sessao-servico.test.js', 'Gerenciamento de sessao'],
        ['Frontend', 'ui-servico.test.js', 'Formatacao e interface'],
        ['Frontend', 'eventos-utils.test.js', 'Utilidades de eventos'],
        ['Frontend', 'geolocalizacao-utils.test.js', 'Calculo de distancias'],
        ['Frontend', 'linha-do-tempo-utils.test.js', 'Utilidades da timeline'],
        ['Frontend', 'pagamentos-utils.test.js', 'Utilidades financeiras'],
    ]
)

h2('6.3  Paginas Impactadas')
body(
    'Os testes cobrem servicos consumidos por todas as paginas do sistema: 3 paginas de autenticacao, '
    '7 paginas do painel da igreja e 7 paginas do painel do membro.',
    indent=True
)

# ============================================================
# 7 SPRINT 6 (NOVO!)
# ============================================================
doc.add_page_break()
h1('7  SPRINT 6 - API GATEWAY, CAPACITACAO E PAGAMENTOS')
body('Periodo: 31 de marco a 06 de abril de 2026.', italic=True)
body(
    'A Sprint 6 focou em tres frentes: implementacao de uma API Gateway do tipo Edge para '
    'centralizar o roteamento de requisicoes, capacitacao de integrantes da equipe em funcionalidades '
    'existentes e inicio do sistema de pagamentos.',
    indent=True
)

# --- 7.1 API Gateway ---
h2('7.1  API Gateway Edge')
body(
    'Joao Pedro e Gabriel implementaram uma API Gateway do tipo Edge que centraliza o roteamento '
    'das requisicoes entre o frontend e o backend. Essa camada intermediaria atua como ponto unico de '
    'entrada para todas as chamadas da aplicacao, direcionando cada requisicao ao servico correto '
    '(Express.js ou FastAPI) de forma transparente para o cliente.',
    indent=True
)
body(
    'A arquitetura Edge executa a logica de roteamento o mais proximo possivel do usuario final, '
    'reduzindo latencia e permitindo tratamento de requisicoes antes que cheguem ao servidor de origem. '
    'Isso possibilita funcionalidades como balanceamento de carga, validacao de autenticacao na borda '
    'e reescrita de rotas.',
    indent=True
)

h2('7.1.1  Beneficios da API Gateway Edge')
bullet('Ponto unico de entrada: todas as requisicoes passam por um unico endpoint, simplificando a configuracao do frontend')
bullet('Roteamento centralizado: regras de redirecionamento entre APIs Express e FastAPI ficam em um unico local')
bullet('Baixa latencia: execucao na borda (Edge) reduz o tempo de resposta para o usuario final')
bullet('Escalabilidade: facilita a adicao de novos servicos backend sem alterar o frontend')
bullet('Seguranca: permite validacao de tokens e rate limiting antes de atingir os servidores de origem')

h2('7.1.2  Arquitetura de Microservicos')
body(
    'A API Gateway roteia para 7 microservicos independentes, cada um responsavel por um dominio '
    'especifico do sistema:',
    indent=True
)
tabela(
    ['Microservico', 'Porta', 'Rota', 'Responsabilidade'],
    [
        ['auth-service', '4001', '/api/auth', 'Autenticacao e autorizacao'],
        ['churches-service', '4002', '/api/igrejas', 'Gestao de igrejas'],
        ['members-service', '4003', '/api/membros', 'Gestao de membros'],
        ['events-service', '4004', '/api/eventos', 'Gestao de eventos'],
        ['finance-service', '4005', '/api/contribuicoes', 'Gestao financeira'],
        ['announcements-service', '4006', '/api/comunicados', 'Comunicados'],
        ['prayers-service', '4007', '/api/pedidos-oracao', 'Pedidos de oracao'],
    ]
)

h2('7.1.3  Middlewares Implementados')
body(
    'A camada de gateway incorpora middlewares que garantem seguranca, observabilidade e '
    'resiliencia:',
    indent=True
)
tabela(
    ['Middleware', 'Funcao'],
    [
        ['request-id', 'Gera identificador unico por requisicao para rastreabilidade'],
        ['logger', 'Registra todas as requisicoes com metodo, rota e tempo de resposta'],
        ['cors', 'Controle de origens permitidas (configuravel por ambiente)'],
        ['express-rate-limit', 'Limite global (200 req/15min) e especifico para auth (20 req/15min)'],
        ['verificarJwt', 'Verificacao de tokens JWT nas rotas protegidas'],
        ['circuit-breaker', 'Padrao Circuit Breaker para resiliencia contra falhas de servicos'],
    ]
)

h2('7.1.4  Fluxo de Requisicoes')
body('O fluxo de uma requisicao no sistema com a API Gateway segue as etapas:', indent=True)
bullet('1. O frontend envia a requisicao para o endpoint da Gateway')
bullet('2. A Gateway recebe, gera request-id e registra no logger')
bullet('3. O rate limiter verifica se o IP nao excedeu o limite de requisicoes')
bullet('4. Para rotas protegidas, o middleware JWT valida o token de autenticacao')
bullet('5. O circuit breaker verifica a saude do servico de destino')
bullet('6. A requisicao e encaminhada ao microservico apropriado')
bullet('7. A Gateway repassa a resposta ao frontend')

# --- 7.2 Capacitacao ---
h2('7.2  Capacitacao da Equipe')
body(
    'Fernando conduziu sessoes de capacitacao com Jhenniffer e Catieli, ensinando na pratica como '
    'funcionam duas funcionalidades implementadas em sprints anteriores: a Linha do Tempo de Eventos '
    'e o QR Code para convite de membros.',
    indent=True
)

h2('7.2.1  Linha do Tempo de Eventos')
body('Foi apresentado o funcionamento completo da pagina de linha do tempo, incluindo:', indent=True)
bullet('Estrutura HTML da visualizacao cronologica vertical')
bullet('Organizacao dos eventos por mes com separadores visuais')
bullet('Contagem regressiva para o proximo evento')
bullet('Filtros por tipo de evento (Culto, Estudo Biblico, Conferencia, etc.)')
bullet('Animacoes de entrada com Intersection Observer')
bullet('Categorizacao por cores conforme o tipo do evento')

h2('7.2.2  QR Code para Convite de Membros')
body('Foi demonstrado como o sistema gera codigos QR para facilitar o cadastro de novos membros:', indent=True)
bullet('Geracao do QR Code no painel da igreja usando qrcode-generator')
bullet('Link de convite com parametro da igreja pre-selecionada')
bullet('Opcoes de compartilhamento: exibir QR, copiar link ou baixar PNG')
bullet('Fluxo completo: pastor gera QR -> membro escaneia -> cadastro com igreja pre-selecionada')

h2('7.2.3  Objetivo da Capacitacao')
body(
    'A capacitacao teve como objetivo nivelar o conhecimento tecnico da equipe, garantindo que todos '
    'os integrantes compreendam as funcionalidades do sistema e possam contribuir em manutencoes '
    'futuras. Jhenniffer e Catieli puderam acompanhar o codigo-fonte, entender a logica de cada '
    'componente e tirar duvidas diretamente com o desenvolvedor responsavel.',
    indent=True
)

# --- 7.3 Pagamentos ---
h2('7.3  Sistema de Pagamentos')
body(
    'Fernando, Catieli e Jhenniffer iniciaram o desenvolvimento do sistema de pagamentos, que permite '
    'o registro e acompanhamento de contribuicoes financeiras (dizimos, ofertas e outros) da igreja.',
    indent=True
)

h2('7.3.1  Funcionalidades Implementadas')
body('Painel da Igreja (administrador):', bold=True)
bullet('Formulario para registrar pagamentos (nome do membro, tipo, valor, data, descricao)')
bullet('Tres tipos de contribuicao: Dizimo, Oferta e Outro')
bullet('Cards de resumo com totais mensais por tipo de contribuicao')
bullet('Tabela com historico completo de pagamentos')
bullet('Opcao de exclusao de registros')

body('Painel do Membro (visualizacao):', bold=True)
bullet('Visualizacao do historico pessoal de contribuicoes')
bullet('Total contribuido e quantidade de registros')
bullet('Layout responsivo: tabela no desktop e cards no mobile')

h2('7.3.2  Bugs Conhecidos')
body(
    'O sistema de pagamentos encontra-se em fase inicial e possui bugs identificados para correcao '
    'nas proximas sprints:',
    indent=True
)
bullet('Filtragem por nome (texto) ao inves de ID unico do membro')
bullet('Sem validacao se membro existe ao registrar pagamento')
bullet('Pagamentos nao editaveis, apenas excluiveis')
bullet('Sem verificacao de lancamentos duplicados')

# --- 7.4 Distribuicao ---
h2('7.4  Distribuicao de Atividades')
tabela(
    ['Integrante', 'Atividades na Sprint 6'],
    [
        ['Fernando Alves da Nobrega',
         'Capacitacao da equipe (linha do tempo e QR Code); Desenvolvimento do sistema de pagamentos'],
        ['Gabriel Franklin Barcellos',
         'Implementacao da API Gateway Edge; Integracao com APIs existentes'],
        ['Joao Pedro Aranda',
         'Arquitetura e implementacao da API Gateway Edge'],
        ['Jhenniffer Lopes da Silva Vargas',
         'Capacitacao em linha do tempo e QR Code; Desenvolvimento do sistema de pagamentos'],
        ['Catieli Gama Cora',
         'Capacitacao em linha do tempo e QR Code; Desenvolvimento do sistema de pagamentos'],
    ]
)

# ============================================================
# 8 ESTRUTURA DO PROJETO
# ============================================================
doc.add_page_break()
h1('8  ESTRUTURA DO PROJETO')
body(
    'Com a evolucao ao longo das seis sprints, o projeto atingiu a seguinte estrutura de diretorios:',
    indent=True
)

tabela(
    ['Pasta', 'Descricao'],
    [
        ['public/', 'Frontend completo (HTML, CSS, JS)'],
        ['public/autenticacao/', 'Login, cadastro com mapa, recuperacao de senha'],
        ['public/igreja/', 'Painel administrativo (8 paginas)'],
        ['public/membros/', 'Painel dos membros (8 paginas)'],
        ['public/js/servicos/', 'Servicos compartilhados (sessao, API, interface)'],
        ['public/js/utils/', 'Utilidades (eventos, geolocalizacao, linha do tempo, pagamentos)'],
        ['api-express/', 'API Gateway com Express.js + middlewares'],
        ['api-express/middlewares/', 'Circuit breaker, JWT, logger, request-id'],
        ['api-fastapi/', 'API REST com FastAPI + Supabase'],
        ['microservices/api-gateway/', 'API Gateway Edge (ponto unico de entrada)'],
        ['microservices/services/', '7 microservicos independentes'],
        ['database/', 'Schema SQL e migracoes'],
        ['tests/', 'Testes unitarios (Vitest) - 9 arquivos de teste'],
    ]
)

body(
    'As paginas do painel da igreja incluem: Painel, Fieis, Eventos, Contribuicoes, Comunicados, '
    'Pedidos de Oracao, Linha do Tempo e configuracoes com QR Code. O painel de membros inclui: '
    'Painel, Eventos, Contribuicoes, Comunicados, Pedidos de Oracao, Linha do Tempo e Meu Perfil.',
    indent=True
)

# ============================================================
# 9 TECNOLOGIAS UTILIZADAS
# ============================================================
doc.add_page_break()
h1('9  TECNOLOGIAS UTILIZADAS')
body(
    'O projeto utiliza exclusivamente tecnologias gratuitas e de codigo aberto:',
    indent=True
)

tabela(
    ['Tecnologia', 'Categoria', 'Finalidade'],
    [
        ['HTML5, CSS3, JavaScript', 'Frontend', 'Interface do usuario (sem frameworks)'],
        ['Express.js', 'Backend', 'API REST e API Gateway (Node.js)'],
        ['FastAPI', 'Backend', 'API REST com validacao automatica (Python)'],
        ['Supabase (PostgreSQL)', 'Banco de Dados', 'Armazenamento com RLS e Auth'],
        ['Firebase Hosting', 'Hospedagem', 'Deploy do frontend com HTTPS'],
        ['Vercel', 'Hospedagem', 'Deploy das APIs (serverless)'],
        ['Vitest', 'Testes', 'Framework de testes unitarios'],
        ['OpenStreetMap + Leaflet.js', 'Web Service', 'Mapa interativo'],
        ['Geolocation API', 'API do navegador', 'Posicao geografica do usuario'],
        ['qrcode-generator', 'Biblioteca JS', 'Geracao de QR Code'],
        ['Intersection Observer', 'API do navegador', 'Animacoes de scroll'],
        ['http-proxy-middleware', 'Infraestrutura', 'Proxy reverso na Gateway'],
        ['express-rate-limit', 'Seguranca', 'Limitacao de requisicoes'],
        ['jsonwebtoken', 'Seguranca', 'Verificacao de tokens JWT'],
    ]
)

# ============================================================
# 10 CONSIDERACOES FINAIS
# ============================================================
doc.add_page_break()
h1('10  CONSIDERACOES FINAIS')
body(
    'O Congrega Fiel evoluiu de um MVP local para um sistema web completo em seis sprints. O '
    'projeto demonstra aplicacao pratica de conceitos como arquitetura REST, consumo de Web Services, '
    'deploy em nuvem, testes automatizados e arquitetura de microservicos, utilizando exclusivamente '
    'tecnologias gratuitas e de codigo aberto.',
    indent=True
)
body(
    'Na Sprint 1, foram definidos o escopo e a arquitetura do sistema. A Sprint 2 migrou o projeto '
    'para a nuvem com banco de dados real e duas APIs REST. A Sprint 3 enriqueceu a experiencia do '
    'usuario com mapa interativo, QR Code e linha do tempo. A Sprint 4 consolidou e refinou essas '
    'funcionalidades. A Sprint 5 implementou testes unitarios com cobertura de backend e frontend.',
    indent=True
)
body(
    'A Sprint 6 representou um avanco significativo em tres frentes: arquitetura, capacitacao e novas '
    'funcionalidades. A implementacao da API Gateway Edge por Joao Pedro e Gabriel moderniza a '
    'comunicacao entre frontend e backend, criando um ponto centralizado de roteamento com '
    'middlewares de seguranca (JWT, rate limiting) e resiliencia (circuit breaker). A capacitacao conduzida '
    'por Fernando garantiu que Jhenniffer e Catieli tenham dominio sobre funcionalidades criticas do '
    'sistema, fortalecendo a autonomia da equipe. O sistema de pagamentos, embora em fase inicial, '
    'representa o primeiro passo para o controle financeiro completo das igrejas.',
    indent=True
)
body(
    'Como proximos passos: correcao dos bugs do sistema de pagamentos, relatorios financeiros, '
    'integracao com gateway de pagamento, aplicativo mobile, notificacoes push, autenticacao na borda '
    'via API Gateway e ampliacao da cobertura de testes.',
    indent=True
)

# ============================================================
# REFERENCIAS
# ============================================================
doc.add_page_break()
h1('REFERENCIAS')

refs = [
    'ABNT. NBR 6023: referencias - elaboracao. Rio de Janeiro, 2018.',
    'ABNT. NBR 14724: trabalhos academicos - apresentacao. Rio de Janeiro, 2011.',
    'EXPRESS.JS. Disponivel em: https://expressjs.com/. Acesso em: 06 abr. 2026.',
    'FASTAPI. Disponivel em: https://fastapi.tiangolo.com/. Acesso em: 06 abr. 2026.',
    'FIREBASE. Disponivel em: https://firebase.google.com/docs/hosting. Acesso em: 06 abr. 2026.',
    'LEAFLET. Disponivel em: https://leafletjs.com/. Acesso em: 06 abr. 2026.',
    'MDN WEB DOCS. Geolocation API. Disponivel em: https://developer.mozilla.org/en-US/docs/Web/API/Geolocation_API. Acesso em: 06 abr. 2026.',
    'MDN WEB DOCS. Intersection Observer API. Disponivel em: https://developer.mozilla.org/en-US/docs/Web/API/Intersection_Observer_API. Acesso em: 06 abr. 2026.',
    'OPENSTREETMAP. Disponivel em: https://www.openstreetmap.org/. Acesso em: 06 abr. 2026.',
    'POSTGRESQL. Disponivel em: https://www.postgresql.org/docs/. Acesso em: 06 abr. 2026.',
    'PRESSMAN, R. S.; MAXIM, B. R. Engenharia de software. 9. ed. Porto Alegre: AMGH, 2021.',
    'SOMMERVILLE, I. Engenharia de software. 10. ed. Sao Paulo: Pearson, 2019.',
    'SUPABASE. Disponivel em: https://supabase.com/docs. Acesso em: 06 abr. 2026.',
    'VERCEL. Disponivel em: https://vercel.com/docs. Acesso em: 06 abr. 2026.',
    'VITEST. Disponivel em: https://vitest.dev. Acesso em: 06 abr. 2026.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(ref)
    run.font.name = FONTE
    run.font.size = Pt(12)

# ============================================================
# SALVAR
# ============================================================
arquivo = 'CongregaFiel - Documentacao Final.docx'
doc.save(arquivo)
print(f'Salvo: {arquivo}')
print(f'{len(doc.paragraphs)} paragrafos, {len(doc.tables)} tabelas')
