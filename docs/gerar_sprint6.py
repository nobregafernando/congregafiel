"""
Gera documento da Sprint 6 - CongregaFiel
API Gateway Edge, Capacitacao da Equipe e Sistema de Pagamentos
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

FONTE = 'Times New Roman'
TAM = Pt(11)


def quebra(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def p(doc, texto, tam=None, bold=False, italic=False, align=None, indent=None):
    par = doc.add_paragraph()
    if align: par.alignment = align
    if indent: par.paragraph_format.first_line_indent = indent
    r = par.add_run(texto)
    r.font.name = FONTE
    r.font.size = tam or TAM
    r.bold = bold
    r.italic = italic
    return par


def h1(doc, texto):
    h = doc.add_heading(texto, level=1)
    for r in h.runs:
        r.font.name = FONTE
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(13)
    return h


def h2(doc, texto):
    h = doc.add_heading(texto, level=2)
    for r in h.runs:
        r.font.name = FONTE
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(11)
    return h


def bullet(doc, texto):
    par = doc.add_paragraph(style='List Bullet')
    r = par.add_run(texto)
    r.font.name = FONTE
    r.font.size = TAM
    return par


def tabela(doc, cabecalho, linhas):
    t = doc.add_table(rows=1 + len(linhas), cols=len(cabecalho))
    try:
        t.style = 'Table Grid'
    except:
        pass
    for j, c in enumerate(cabecalho):
        cell = t.rows[0].cells[j]
        cell.text = c
        for par in cell.paragraphs:
            for r in par.runs:
                r.bold = True
                r.font.name = FONTE
                r.font.size = Pt(10)
    for i, linha in enumerate(linhas):
        for j, val in enumerate(linha):
            cell = t.rows[i + 1].cells[j]
            cell.text = val
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.name = FONTE
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return t


# ============================================================
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = FONTE
style.font.size = TAM
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

for i in range(1, 4):
    s = doc.styles[f'Heading {i}']
    s.font.name = FONTE
    s.font.color.rgb = RGBColor(0, 0, 0)

# ============================================================
# CAPA
# ============================================================
for _ in range(3):
    doc.add_paragraph()
p(doc, 'FACULDADE INSTED', Pt(14), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, 'Curso Superior de Tecnologia em Analise e Desenvolvimento de Sistemas',
  Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(4):
    doc.add_paragraph()
p(doc, 'CONGREGA FIEL', Pt(16), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, 'Sistema Web para Gestao de Comunidades Eclesiasticas', Pt(12),
  align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(2):
    doc.add_paragraph()
p(doc, 'Sprint 6 - API Gateway, Capacitacao e Pagamentos', Pt(12), bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(5):
    doc.add_paragraph()
for nome in ['Catieli Gama Cora', 'Fernando Alves da Nobrega',
             'Gabriel Franklin Barcellos', 'Jhenniffer Lopes da Silva Vargas',
             'Joao Pedro Aranda']:
    p(doc, nome, Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(4):
    doc.add_paragraph()
p(doc, 'Campo Grande - MS', Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, '2026', Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# 1. INTRODUCAO
# ============================================================
quebra(doc)
h1(doc, '1  INTRODUCAO')

p(doc, 'Este documento registra as atividades desenvolvidas na Sprint 6 do projeto '
  'Congrega Fiel, correspondente ao periodo de 31 de marco a 06 de abril de 2026. '
  'Nesta sprint, o foco foi a implementacao de uma API Gateway do tipo Edge para '
  'centralizar o roteamento de requisicoes, a capacitacao de integrantes da equipe '
  'em funcionalidades ja existentes e o inicio do sistema de pagamentos.',
  indent=Cm(1.25))

p(doc, 'Periodo: 31 de marco a 06 de abril de 2026.', italic=True)

# ============================================================
# 2. ATIVIDADES DA SPRINT
# ============================================================
quebra(doc)
h1(doc, '2  ATIVIDADES REALIZADAS')

# --- 2.1 API Gateway ---
h2(doc, '2.1  API Gateway Edge')

p(doc, 'Joao Pedro e Gabriel implementaram uma API Gateway do tipo Edge que centraliza '
  'o roteamento das requisicoes entre o frontend e o backend. Essa camada intermediaria '
  'atua como ponto unico de entrada para todas as chamadas da aplicacao, direcionando '
  'cada requisicao ao servico correto (Express.js ou FastAPI) de forma transparente '
  'para o cliente.',
  indent=Cm(1.25))

p(doc, 'A arquitetura Edge executa a logica de roteamento o mais proximo possivel do '
  'usuario final, reduzindo latencia e permitindo tratamento de requisicoes antes '
  'que cheguem ao servidor de origem. Isso possibilita funcionalidades como '
  'balanceamento de carga, validacao de autenticacao na borda e reescrita de rotas.',
  indent=Cm(1.25))

h2(doc, '2.1.1  Beneficios da API Gateway Edge')
bullet(doc, 'Ponto unico de entrada: todas as requisicoes passam por um unico endpoint, '
       'simplificando a configuracao do frontend')
bullet(doc, 'Roteamento centralizado: regras de redirecionamento entre APIs Express e '
       'FastAPI ficam em um unico local')
bullet(doc, 'Baixa latencia: execucao na borda (Edge) reduz o tempo de resposta para '
       'o usuario final')
bullet(doc, 'Escalabilidade: facilita a adicao de novos servicos backend sem alterar '
       'o frontend')
bullet(doc, 'Seguranca: permite validacao de tokens e rate limiting antes de atingir '
       'os servidores de origem')

h2(doc, '2.1.2  Fluxo de Requisicoes')
p(doc, 'O fluxo de uma requisicao no sistema com a API Gateway segue as etapas:',
  indent=Cm(1.25))
bullet(doc, '1. O frontend envia a requisicao para o endpoint da Gateway')
bullet(doc, '2. A Gateway recebe e analisa a rota solicitada')
bullet(doc, '3. Com base nas regras de roteamento, a requisicao e encaminhada '
       'ao backend apropriado (Express ou FastAPI)')
bullet(doc, '4. O backend processa a requisicao e retorna a resposta')
bullet(doc, '5. A Gateway repassa a resposta ao frontend')

h2(doc, '2.1.3  Responsaveis')
tabela(doc, ['Integrante', 'Contribuicao'], [
    ['Joao Pedro Aranda', 'Arquitetura e implementacao da Gateway Edge'],
    ['Gabriel Franklin Barcellos', 'Implementacao e integracao com APIs existentes'],
])

# --- 2.2 Capacitacao ---
quebra(doc)
h2(doc, '2.2  Capacitacao da Equipe')

p(doc, 'Fernando conduziu sessoes de capacitacao com Jhenniffer e Catieli, ensinando '
  'na pratica como funcionam duas funcionalidades implementadas em sprints anteriores: '
  'a Linha do Tempo de Eventos e o QR Code para convite de membros.',
  indent=Cm(1.25))

h2(doc, '2.2.1  Linha do Tempo de Eventos')
p(doc, 'Foi apresentado o funcionamento completo da pagina de linha do tempo, incluindo:',
  indent=Cm(1.25))
bullet(doc, 'Estrutura HTML da visualizacao cronologica vertical')
bullet(doc, 'Organizacao dos eventos por mes com separadores visuais')
bullet(doc, 'Contagem regressiva para o proximo evento')
bullet(doc, 'Filtros por tipo de evento (Culto, Estudo Biblico, Conferencia, etc.)')
bullet(doc, 'Animacoes de entrada com Intersection Observer')
bullet(doc, 'Categorizacao por cores conforme o tipo do evento')

h2(doc, '2.2.2  QR Code para Convite de Membros')
p(doc, 'Foi demonstrado como o sistema gera codigos QR para facilitar o cadastro '
  'de novos membros na igreja:',
  indent=Cm(1.25))
bullet(doc, 'Geracao do QR Code no painel da igreja usando qrcode-generator')
bullet(doc, 'Link de convite com parametro da igreja pre-selecionada')
bullet(doc, 'Opcoes de compartilhamento: exibir QR, copiar link ou baixar PNG')
bullet(doc, 'Fluxo completo: pastor gera QR -> membro escaneia -> cadastro com igreja pre-selecionada')

h2(doc, '2.2.3  Objetivo da Capacitacao')
p(doc, 'A capacitacao teve como objetivo nivelar o conhecimento tecnico da equipe, '
  'garantindo que todos os integrantes compreendam as funcionalidades do sistema '
  'e possam contribuir em manutencoes futuras. Jhenniffer e Catieli puderam '
  'acompanhar o codigo-fonte, entender a logica de cada componente e tirar duvidas '
  'diretamente com o desenvolvedor responsavel.',
  indent=Cm(1.25))

h2(doc, '2.2.4  Responsaveis')
tabela(doc, ['Integrante', 'Papel'], [
    ['Fernando Alves da Nobrega', 'Instrutor (ensinou as funcionalidades)'],
    ['Jhenniffer Lopes da Silva Vargas', 'Aprendiz (linha do tempo e QR Code)'],
    ['Catieli Gama Cora', 'Aprendiz (linha do tempo e QR Code)'],
])

# --- 2.3 Pagamentos ---
quebra(doc)
h2(doc, '2.3  Sistema de Pagamentos')

p(doc, 'Fernando, Catieli e Jhenniffer iniciaram o desenvolvimento do sistema de '
  'pagamentos, que permite o registro e acompanhamento de contribuicoes financeiras '
  '(dizimos, ofertas e outros) da igreja. Nesta versao inicial, os dados sao '
  'armazenados via banco de dados local.',
  indent=Cm(1.25))

h2(doc, '2.3.1  Funcionalidades Implementadas')
p(doc, 'Painel da Igreja (administrador):', bold=True)
bullet(doc, 'Formulario para registrar pagamentos (nome do membro, tipo, valor, data, descricao)')
bullet(doc, 'Tres tipos de contribuicao: Dizimo, Oferta e Outro')
bullet(doc, 'Cards de resumo com totais mensais por tipo de contribuicao')
bullet(doc, 'Tabela com historico completo de pagamentos')
bullet(doc, 'Opcao de exclusao de registros')

doc.add_paragraph()
p(doc, 'Painel do Membro (visualizacao):', bold=True)
bullet(doc, 'Visualizacao do historico pessoal de contribuicoes')
bullet(doc, 'Total contribuido e quantidade de registros')
bullet(doc, 'Layout responsivo: tabela no desktop e cards no mobile')

h2(doc, '2.3.2  Status Atual e Bugs Conhecidos')
p(doc, 'O sistema de pagamentos encontra-se em fase inicial e ainda possui bugs que '
  'serao corrigidos nas proximas sprints. Os principais problemas identificados sao:',
  indent=Cm(1.25))

bullet(doc, 'Filtragem por nome: a identificacao do membro na visualizacao de pagamentos '
       'e feita por comparacao de texto (nome), e nao por ID unico, o que pode gerar '
       'inconsistencias caso existam membros com nomes iguais')
bullet(doc, 'Validacao de membro: o administrador pode registrar pagamentos para nomes '
       'que nao correspondem a membros cadastrados na igreja')
bullet(doc, 'Ausencia de edicao: pagamentos registrados so podem ser excluidos, nao '
       'editados, dificultando correcoes de lancamentos errados')
bullet(doc, 'Precisao decimal: uso de ponto flutuante do JavaScript pode causar '
       'pequenas imprecisoes em valores monetarios')
bullet(doc, 'Duplicidade: nao ha verificacao para evitar lancamentos duplicados')

h2(doc, '2.3.3  Responsaveis')
tabela(doc, ['Integrante', 'Contribuicao'], [
    ['Fernando Alves da Nobrega', 'Arquitetura e desenvolvimento do modulo'],
    ['Catieli Gama Cora', 'Desenvolvimento e testes do formulario'],
    ['Jhenniffer Lopes da Silva Vargas', 'Desenvolvimento e layout responsivo'],
])

# ============================================================
# 3. DISTRIBUICAO DE ATIVIDADES
# ============================================================
quebra(doc)
h1(doc, '3  DISTRIBUICAO DE ATIVIDADES')

tabela(doc, ['Integrante', 'Atividades na Sprint 6'], [
    ['Fernando Alves da Nobrega',
     'Capacitacao da equipe (linha do tempo e QR Code); '
     'Desenvolvimento do sistema de pagamentos'],
    ['Gabriel Franklin Barcellos',
     'Implementacao da API Gateway Edge; '
     'Integracao com APIs existentes'],
    ['Joao Pedro Aranda',
     'Arquitetura e implementacao da API Gateway Edge'],
    ['Jhenniffer Lopes da Silva Vargas',
     'Capacitacao em linha do tempo e QR Code; '
     'Desenvolvimento do sistema de pagamentos'],
    ['Catieli Gama Cora',
     'Capacitacao em linha do tempo e QR Code; '
     'Desenvolvimento do sistema de pagamentos'],
])

# ============================================================
# 4. TECNOLOGIAS UTILIZADAS
# ============================================================
quebra(doc)
h1(doc, '4  TECNOLOGIAS UTILIZADAS NA SPRINT')

tabela(doc, ['Tecnologia', 'Finalidade'], [
    ['Edge Functions', 'Execucao da API Gateway na borda (Edge)'],
    ['Vercel / Supabase Edge', 'Hospedagem das funcoes de roteamento'],
    ['Express.js', 'Backend Node.js (API de destino)'],
    ['FastAPI', 'Backend Python (API de destino)'],
    ['Supabase (PostgreSQL)', 'Armazenamento dos dados de pagamentos'],
    ['HTML/CSS/JavaScript', 'Interface do modulo de pagamentos'],
    ['Leaflet.js', 'Mapa interativo (objeto de capacitacao)'],
    ['qrcode-generator', 'Geracao de QR Code (objeto de capacitacao)'],
])

# ============================================================
# 5. PROXIMOS PASSOS
# ============================================================
quebra(doc)
h1(doc, '5  PROXIMOS PASSOS')

bullet(doc, 'Corrigir os bugs do sistema de pagamentos (filtragem por ID, validacao de membro, edicao de registros)')
bullet(doc, 'Adicionar relatorios financeiros (mensal, anual, por membro)')
bullet(doc, 'Implementar trilha de auditoria para operacoes financeiras')
bullet(doc, 'Expandir a API Gateway com autenticacao na borda e rate limiting')
bullet(doc, 'Testes unitarios para o modulo de pagamentos e a API Gateway')

# ============================================================
# 6. CONSIDERACOES FINAIS
# ============================================================
quebra(doc)
h1(doc, '6  CONSIDERACOES FINAIS')

p(doc, 'A Sprint 6 representou um avanco significativo em tres frentes: arquitetura, '
  'capacitacao e novas funcionalidades. A implementacao da API Gateway Edge por Joao e '
  'Gabriel moderniza a comunicacao entre frontend e backend, criando um ponto centralizado '
  'de roteamento que facilita a manutencao e escalabilidade do sistema.',
  indent=Cm(1.25))

p(doc, 'A capacitacao conduzida por Fernando garantiu que Jhenniffer e Catieli tenham '
  'dominio sobre funcionalidades criticas do sistema, fortalecendo a autonomia da equipe '
  'como um todo. Esse nivelamento tecnico e fundamental para a continuidade do projeto.',
  indent=Cm(1.25))

p(doc, 'O sistema de pagamentos, embora ainda em fase inicial com bugs conhecidos, '
  'representa o primeiro passo para o controle financeiro completo das igrejas. '
  'As correcoes e melhorias planejadas para as proximas sprints tornarao o modulo '
  'robusto e confiavel para uso em producao.',
  indent=Cm(1.25))

# SALVAR
arquivo = 'Semana 6 - Congregafiel.docx'
doc.save(arquivo)
print(f"Salvo: {arquivo}")
print(f"{len(doc.paragraphs)} paragrafos, {len(doc.tables)} tabelas")
