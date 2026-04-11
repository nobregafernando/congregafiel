"""
Gera documento unificado resumido do CongregaFiel
Todas as sprints em um unico fluxo, sem capa, sem quebras de pagina
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONTE = 'Times New Roman'
TAM = Pt(11)


def p(doc, texto, tam=None, bold=False, italic=False, indent=None):
    par = doc.add_paragraph()
    if indent: par.paragraph_format.first_line_indent = indent
    r = par.add_run(texto)
    r.font.name = FONTE
    r.font.size = tam or TAM
    r.bold = bold
    r.italic = italic
    return par


def h1(doc, texto):
    h = doc.add_heading(texto, level=1)
    for r in h.runs:
        r.font.name = FONTE
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(13)
    return h


def h2(doc, texto):
    h = doc.add_heading(texto, level=2)
    for r in h.runs:
        r.font.name = FONTE
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(11)
    return h


def bullet(doc, texto):
    par = doc.add_paragraph(style='List Bullet')
    r = par.add_run(texto)
    r.font.name = FONTE
    r.font.size = Pt(10)
    par.paragraph_format.space_after = Pt(1)
    return par


def tabela(doc, cabecalho, linhas):
    t = doc.add_table(rows=1 + len(linhas), cols=len(cabecalho))
    try:
        t.style = 'Table Grid'
    except:
        pass
    for j, c in enumerate(cabecalho):
        cell = t.rows[0].cells[j]
        cell.text = c
        for par in cell.paragraphs:
            for r in par.runs:
                r.bold = True
                r.font.name = FONTE
                r.font.size = Pt(9)
    for i, linha in enumerate(linhas):
        for j, val in enumerate(linha):
            cell = t.rows[i + 1].cells[j]
            cell.text = val
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.name = FONTE
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return t


# ============================================================
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = FONTE
style.font.size = TAM
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(2)

for i in range(1, 4):
    s = doc.styles[f'Heading {i}']
    s.font.name = FONTE
    s.font.color.rgb = RGBColor(0, 0, 0)

# ============================================================
# TITULO
# ============================================================
p(doc, 'CONGREGA FIEL - Registro de Sprints', Pt(14), bold=True)
p(doc, 'Sistema Web para Gestao de Comunidades Eclesiasticas', Pt(10), italic=True)
p(doc, 'Equipe: Catieli, Fernando, Gabriel, Jhenniffer e Joao Pedro', Pt(10))
p(doc, 'Periodo: 24/02/2026 a 06/04/2026 (6 sprints)', Pt(10))

# ============================================================
# SPRINT 1
# ============================================================
h1(doc, 'Sprint 1 - Documentacao do MVP (24/02 a 02/03)')

p(doc, 'Definicao do escopo, requisitos e arquitetura do sistema. Problema: igrejas '
  'de pequeno porte sem ferramenta centralizada para gestao de membros, financas e eventos.',
  indent=Cm(1.25))

p(doc, 'Modulos definidos: Autenticacao, Gestao da Igreja, Membros, Financeiro, Eventos e Comunicados. '
  'Dois perfis de acesso: Pastor (admin) e Fiel (membro). Arquitetura REST com JSON Server '
  'e LocalStorage na fase inicial.',
  indent=Cm(1.25))

# ============================================================
# SPRINT 2
# ============================================================
h1(doc, 'Sprint 2 - Deploy Online e Banco de Dados (03/03 a 09/03)')

bullet(doc, 'Frontend publicado no Firebase Hosting (HTTPS automatico)')
bullet(doc, 'Banco PostgreSQL no Supabase: 6 tabelas com Row Level Security')
bullet(doc, 'API Express.js (33 rotas) e API FastAPI (33 rotas), ambas no Vercel')
bullet(doc, 'Autenticacao com Supabase Auth (JWT + bcrypt)')

tabela(doc, ['Componente', 'URL'], [
    ['Frontend', 'https://congregafiel.web.app'],
    ['API Express', 'https://api-express-tau.vercel.app'],
    ['API FastAPI', 'https://api-fastapi.vercel.app'],
])

# ============================================================
# SPRINT 3
# ============================================================
h1(doc, 'Sprint 3 - Web Services e Melhorias (10/03 a 16/03)')

bullet(doc, 'Mapa interativo (OpenStreetMap + Leaflet.js) no cadastro de membros com geolocalizacao')
bullet(doc, 'QR Code para convite de membros (qrcode-generator via CDN)')
bullet(doc, 'Linha do tempo de eventos com filtros por tipo e animacoes (Intersection Observer)')
bullet(doc, 'Categorizacao de eventos por tipo e cor (Culto, Estudo, Conferencia, Especial, Outro)')
bullet(doc, 'Carrossel de eventos, perfil do membro, filtros em pedidos de oracao')

# ============================================================
# SPRINT 4
# ============================================================
h1(doc, 'Sprint 4 - Consolidacao e Refinamento (17/03 a 23/03)')

bullet(doc, 'Responsividade do mapa interativo para dispositivos moveis')
bullet(doc, 'Otimizacao das animacoes da linha do tempo')
bullet(doc, 'Validacao do fluxo completo de convite via QR Code')
bullet(doc, 'Sistema de respostas nos pedidos de oracao')
bullet(doc, 'Ajustes na pagina de perfil do membro')

# ============================================================
# SPRINT 5
# ============================================================
h1(doc, 'Sprint 5 - Testes Unitarios (24/03 a 30/03)')

p(doc, 'Implementacao de testes unitarios com Vitest cobrindo backend e frontend:',
  indent=Cm(1.25))

tabela(doc, ['Camada', 'Arquivo de Teste', 'Cobertura'], [
    ['Backend', 'pedidos-oracao-utils.test.js', 'Pedidos de oracao'],
    ['Backend', 'regras-auth.test.js', 'Regras de autenticacao'],
    ['Frontend', 'api-servico.test.js', 'Comunicacao com API'],
    ['Frontend', 'sessao-servico.test.js', 'Gerenciamento de sessao'],
    ['Frontend', 'ui-servico.test.js', 'Formatacao e interface'],
    ['Frontend', 'eventos-utils.test.js', 'Utilidades de eventos'],
    ['Frontend', 'geolocalizacao-utils.test.js', 'Calculo de distancias'],
    ['Frontend', 'linha-do-tempo-utils.test.js', 'Utilidades da timeline'],
    ['Frontend', 'pagamentos-utils.test.js', 'Utilidades financeiras'],
])

# ============================================================
# SPRINT 6
# ============================================================
h1(doc, 'Sprint 6 - API Gateway, Capacitacao e Pagamentos (31/03 a 06/04)')

h2(doc, 'API Gateway Edge')
p(doc, 'Joao Pedro e Gabriel implementaram uma API Gateway do tipo Edge que centraliza '
  'o roteamento das requisicoes entre frontend e backend. Atua como ponto unico de '
  'entrada, direcionando cada chamada ao servico correto (Express ou FastAPI) com '
  'baixa latencia por executar na borda.',
  indent=Cm(1.25))

h2(doc, 'Capacitacao da Equipe')
p(doc, 'Fernando ensinou Jhenniffer e Catieli a montar a Linha do Tempo de Eventos '
  '(estrutura HTML, organizacao por mes, filtros, animacoes) e o QR Code para convite '
  'de membros (geracao, link com igreja pre-selecionada, opcoes de compartilhamento). '
  'Objetivo: nivelar o conhecimento tecnico para manutencoes futuras.',
  indent=Cm(1.25))

h2(doc, 'Sistema de Pagamentos')
p(doc, 'Fernando, Catieli e Jhenniffer iniciaram o modulo de pagamentos (dizimos, ofertas, outros). '
  'Admin registra pagamentos com resumo mensal; membro visualiza historico pessoal com layout responsivo.',
  indent=Cm(1.25))

p(doc, 'Bugs conhecidos:', bold=True)
bullet(doc, 'Filtragem por nome (texto) ao inves de ID unico')
bullet(doc, 'Sem validacao se membro existe ao registrar pagamento')
bullet(doc, 'Pagamentos nao editaveis, apenas excluiveis')
bullet(doc, 'Sem verificacao de lancamentos duplicados')

h2(doc, 'Distribuicao de Atividades')
tabela(doc, ['Integrante', 'Atividades'], [
    ['Fernando', 'Capacitacao (linha do tempo + QR Code); Sistema de pagamentos'],
    ['Gabriel', 'Implementacao da API Gateway Edge; Integracao com APIs'],
    ['Joao Pedro', 'Arquitetura e implementacao da API Gateway Edge'],
    ['Jhenniffer', 'Capacitacao (linha do tempo + QR Code); Sistema de pagamentos'],
    ['Catieli', 'Capacitacao (linha do tempo + QR Code); Sistema de pagamentos'],
])

# SALVAR
arquivo = 'CongregaFiel - Documentacao Completa.docx'
doc.save(arquivo)
print(f"Salvo: {arquivo}")
print(f"{len(doc.paragraphs)} paragrafos, {len(doc.tables)} tabelas")
