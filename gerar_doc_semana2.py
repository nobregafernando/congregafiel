# -*- coding: utf-8 -*-
"""
Documentação Semana 2 — Congrega Fiel (MVP)
Web API: Express.js e FastAPI
Faculdade INSTED · Campo Grande/MS
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import os

# ============================================================
# PALETA COMERCIAL (idêntica à Semana 1)
# ============================================================
CLR_PRIMARY   = "1B2A4A"
CLR_ACCENT    = "C5A55A"
CLR_LIGHT_BG  = "F0EDE6"
CLR_WHITE     = "FFFFFF"
CLR_DARK_TEXT = "1B2A4A"
CLR_MID_GRAY  = "6B7280"
CLR_LIGHT_LN  = "D4C99E"
CLR_TABLE_HDR = "1B2A4A"

# ============================================================
# CONFIG GERAL
# ============================================================
FONT_NAME = "Times New Roman"
SZ = 12
SZ_TITLE = 14
SZ_COVER = 16
SZ_BIG = 28
LSPACING = 1.5

INST  = "FACULDADE INSTED"
CURSO = "Curso Superior de Tecnologia em Análise e Desenvolvimento de Sistemas"
CITY  = "Campo Grande \u2013 MS"
YEAR  = "2026"

doc = Document()

# Margens ABNT
for sec in doc.sections:
    sec.top_margin    = Cm(3)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2)
    sec.page_height   = Cm(29.7)
    sec.page_width    = Cm(21)

# Auto-update fields (sumário)
uf = OxmlElement("w:updateFields")
uf.set(qn("w:val"), "true")
doc.settings.element.append(uf)

# ============================================================
# ESTILOS BASE
# ============================================================
sn = doc.styles["Normal"]
sn.font.name = FONT_NAME
sn.font.size = Pt(SZ)
sn.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
sn.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
sn.paragraph_format.line_spacing = LSPACING
sn.paragraph_format.space_after  = Pt(0)
sn.paragraph_format.space_before = Pt(0)
rPr = sn.element.get_or_add_rPr()
rPr.append(parse_xml(
    f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" '
    f'w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>'
))

for lvl, name in enumerate(["Heading 1", "Heading 2", "Heading 3"], 1):
    hs = doc.styles[name]
    hs.font.name = FONT_NAME
    hs.font.size = Pt(SZ)
    hs.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    hs.font.bold = lvl <= 2
    hs.font.italic = lvl == 3
    hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hs.paragraph_format.line_spacing = LSPACING
    hs.paragraph_format.space_before = Pt(18 if lvl == 1 else 12 if lvl == 2 else 8)
    hs.paragraph_format.space_after  = Pt(6 if lvl == 1 else 4)
    hs.paragraph_format.first_line_indent = Cm(0)
    hs.paragraph_format.left_indent = Cm(0)
    pPr = hs.element.get_or_add_pPr()
    num = pPr.find(qn("w:numPr"))
    if num is not None:
        pPr.remove(num)
    hrPr = hs.element.get_or_add_rPr()
    hrPr.append(parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" '
        f'w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>'
    ))
    hrPr.append(parse_xml(f'<w:color {nsdecls("w")} w:val="{CLR_PRIMARY}"/>'))

# ============================================================
# HELPERS (idênticos à Semana 1)
# ============================================================
def blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(); r.font.size = Pt(SZ); r.font.name = FONT_NAME

def centered(text, size=SZ, bold=False, upper=False, color=CLR_DARK_TEXT):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LSPACING
    r = p.add_run(text.upper() if upper else text)
    r.font.size = Pt(size); r.font.name = FONT_NAME; r.bold = bold
    r.font.color.rgb = RGBColor(int(color[:2],16), int(color[2:4],16), int(color[4:],16))
    return p

def body(text, size=SZ, bold=False, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LSPACING
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = FONT_NAME; r.bold = bold
    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return p

def colored_line(color=CLR_ACCENT, thickness=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{thickness}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def section_h1(number, title):
    text = f"{number}  {title.upper()}" if number else title.upper()
    p = doc.add_paragraph(text, style="Heading 1")
    for r in p.runs:
        r.font.size = Pt(SZ); r.font.name = FONT_NAME; r.bold = True
        r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="2" w:color="{CLR_ACCENT}"/>'
        f'</w:pBdr>'
    ))
    return p

def section_h2(number, title):
    p = doc.add_paragraph(f"{number}  {title}", style="Heading 2")
    for r in p.runs:
        r.font.size = Pt(SZ); r.font.name = FONT_NAME; r.bold = True
        r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return p

def section_h3(number, title):
    p = doc.add_paragraph(f"{number}  {title}", style="Heading 3")
    for r in p.runs:
        r.font.size = Pt(SZ); r.font.name = FONT_NAME
        r.bold = False; r.italic = True
        r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = LSPACING
    ind = 1.25 + (level * 0.75)
    p.paragraph_format.left_indent = Cm(ind)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    sym = "\u2022" if level == 0 else "\u25E6"
    r = p.add_run(f"{sym}  {text}")
    r.font.size = Pt(SZ); r.font.name = FONT_NAME
    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"

    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after  = Pt(4)
        rr = pp.add_run(h)
        rr.font.size = Pt(10); rr.font.name = FONT_NAME; rr.bold = True
        rr.font.color.rgb = RGBColor(0xC5, 0xA5, 0x5A)
        c._element.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CLR_TABLE_HDR}" w:val="clear"/>')
        )

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ""
            pp = c.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pp.paragraph_format.space_before = Pt(3)
            pp.paragraph_format.space_after  = Pt(3)
            rr = pp.add_run(str(val))
            rr.font.size = Pt(10); rr.font.name = FONT_NAME
            rr.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
            if ri % 2 == 0:
                c._element.get_or_add_tcPr().append(
                    parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CLR_WHITE}" w:val="clear"/>')
                )
            else:
                c._element.get_or_add_tcPr().append(
                    parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CLR_LIGHT_BG}" w:val="clear"/>')
                )

    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'  <w:left   w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'  <w:right  w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    blank(1)
    return t

def toc_field():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = LSPACING

    r1 = p.add_run()
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    r1._element.append(fc1)

    r2 = p.add_run()
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
    ins.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._element.append(ins)

    r3 = p.add_run()
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate")
    r3._element.append(fc2)

    r4 = p.add_run("Abra no Word e pressione F9 para gerar o sumário linkado.")
    r4.font.size = Pt(SZ); r4.font.name = FONT_NAME
    r4.font.color.rgb = RGBColor(0x6B, 0x72, 0x80); r4.italic = True

    r5 = p.add_run()
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end")
    r5._element.append(fc3)

def page_break():
    doc.add_page_break()

def cover_band(color, height_cm=1.2):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.text = ""
    cell.height = Cm(height_cm)
    cell._element.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    )
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left   w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right  w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    ))
    tblPr.append(parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    ))

def code_block(text):
    """Bloco de código com fundo creme e fonte monoespaçada."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(1.25)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{CLR_LIGHT_BG}" w:val="clear"/>'
    ))
    r = p.add_run(text)
    r.font.size = Pt(9); r.font.name = "Consolas"
    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return p

def bold_body(label, text):
    """Parágrafo com rótulo em negrito seguido de texto normal."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LSPACING
    p.paragraph_format.first_line_indent = Cm(1.25)
    rb = p.add_run(f"{label}: ")
    rb.font.size = Pt(SZ); rb.font.name = FONT_NAME; rb.bold = True
    rb.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    rd = p.add_run(text)
    rd.font.size = Pt(SZ); rd.font.name = FONT_NAME
    rd.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return p


# ============================================================
#                       CAPA
# ============================================================
blank(1)
cover_band(CLR_PRIMARY, 0.6)
blank(1)
centered(INST, size=SZ_COVER, bold=True, upper=True, color=CLR_PRIMARY)
colored_line(CLR_ACCENT, 6)
centered(CURSO, size=11, bold=False, color=CLR_MID_GRAY)
blank(4)
centered("CONGREGA FIEL", size=SZ_BIG, bold=True, upper=True, color=CLR_PRIMARY)
blank(1)
colored_line(CLR_ACCENT, 18)
blank(1)
centered("Semana 2 \u2014 Web API com Express.js, FastAPI e Supabase", size=SZ_TITLE, bold=False, color=CLR_MID_GRAY)
blank(3)

for nome in [
    "Catieli Gama Cora",
    "Fernando Alves da Nóbrega",
    "Gabriel Franklin Barcellos",
    "Jhenniffer Lopes da Silva Vargas",
    "João Pedro Aranda",
]:
    centered(nome, size=SZ, bold=False, color=CLR_PRIMARY)

blank(4)
cover_band(CLR_ACCENT, 0.15)
blank(1)
centered(CITY, size=SZ, bold=False, color=CLR_MID_GRAY)
centered(YEAR, size=SZ, bold=True, color=CLR_PRIMARY)

# ============================================================
#                    FOLHA DE ROSTO
# ============================================================
page_break()
blank(2)
centered("CONGREGA FIEL", size=22, bold=True, upper=True, color=CLR_PRIMARY)
colored_line(CLR_ACCENT, 10)
centered("Semana 2 \u2014 Web API com Express.js, FastAPI e Supabase",
         size=SZ_TITLE, bold=False, color=CLR_MID_GRAY)
blank(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0
r = p.add_run(
    f"Documentação técnica da Sprint 2 \u2014 implementação de Web APIs REST "
    f"utilizando dois frameworks: Express.js (Node.js) e FastAPI (Python), "
    f"com persistência de dados no Supabase (PostgreSQL). "
    f"Apresentada como requisito parcial para aprovação no {CURSO} da {INST}."
)
r.font.size = Pt(10); r.font.name = FONT_NAME
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

blank(3)
centered("Equipe de Desenvolvimento", size=SZ, bold=True, color=CLR_PRIMARY)
colored_line(CLR_ACCENT, 4)
blank(1)
for nome in [
    "Catieli Gama Cora",
    "Fernando Alves da Nóbrega",
    "Gabriel Franklin Barcellos",
    "Jhenniffer Lopes da Silva Vargas",
    "João Pedro Aranda",
]:
    centered(nome, size=SZ, color=CLR_PRIMARY)

blank(5)
cover_band(CLR_ACCENT, 0.1)
blank(1)
centered(CITY, size=SZ, color=CLR_MID_GRAY)
centered(YEAR, size=SZ, bold=True, color=CLR_PRIMARY)

# ============================================================
#                       SUMÁRIO
# ============================================================
page_break()
centered("SUMÁRIO", size=SZ_TITLE, bold=True, upper=True, color=CLR_PRIMARY)
colored_line(CLR_ACCENT, 8)
blank(1)
toc_field()

# ============================================================
# 1  INTRODUÇÃO
# ============================================================
page_break()
section_h1("1", "INTRODUÇÃO")

body(
    "Este documento apresenta a implementação das Web APIs REST do projeto "
    "Congrega Fiel, desenvolvidas durante a Sprint 2 (03/03 \u2013 09/03/2026). "
    "Conforme definido no cronograma do MVP, esta etapa contempla a construção "
    "da camada de serviços que permite a comunicação entre o front-end e os dados "
    "do sistema, além da migração da persistência de dados para o Supabase."
)

body(
    "Foram implementados dois frameworks de Web API para atender ao requisito "
    "acadêmico de demonstrar domínio em múltiplas tecnologias: Express.js, "
    "baseado em Node.js (JavaScript), e FastAPI, baseado em Python. Ambos "
    "implementam os mesmos endpoints RESTful e se conectam ao Supabase como "
    "banco de dados, permitindo uma comparação direta entre as abordagens."
)

body(
    "A documentação a seguir detalha a arquitetura das APIs, a configuração "
    "do banco de dados Supabase, os endpoints implementados, os modelos de dados, "
    "exemplos de requisições e respostas, uma explicação prática de como cada "
    "framework funciona no sistema, e uma análise comparativa."
)

# ============================================================
# 2  OBJETIVOS DA SPRINT 2
# ============================================================
page_break()
section_h1("2", "OBJETIVOS DA SPRINT 2")

body("Os objetivos desta sprint foram definidos com base no cronograma do MVP:")

bullet("Implementar uma API REST completa com Express.js (Node.js);")
bullet("Implementar uma API REST equivalente com FastAPI (Python);")
bullet("Criar endpoints para todos os módulos do sistema: igrejas, membros, eventos, contribuições, comunicados e pedidos de oração;")
bullet("Configurar o Supabase (PostgreSQL) como banco de dados em nuvem;")
bullet("Migrar a persistência de arquivos JSON para o Supabase;")
bullet("Documentar os endpoints, modelos de dados e exemplos de uso;")
bullet("Comparar os dois frameworks em termos de funcionalidades e características.")

# ============================================================
# 3  FRAMEWORK 1 — EXPRESS.JS
# ============================================================
page_break()
section_h1("3", "FRAMEWORK 1 \u2014 EXPRESS.JS (NODE.JS)")

section_h2("3.1", "Visão Geral")

body(
    "Express.js é o framework web mais popular para Node.js, com mais de 60 mil "
    "estrelas no GitHub. Criado em 2010, é conhecido por sua abordagem minimalista "
    "e flexível, permitindo criar APIs REST de forma rápida e com baixa curva de "
    "aprendizado. Utiliza a mesma linguagem do front-end (JavaScript), o que "
    "facilita a integração em projetos web."
)

section_h2("3.2", "Configuração e Dependências")

body("O projeto Express.js utiliza as seguintes dependências:")

table(
    ["Pacote", "Versão", "Finalidade"],
    [
        ["express", "4.21.0", "Framework web para criação de rotas e middlewares"],
        ["cors", "2.8.5", "Middleware para permitir requisições cross-origin"],
        ["@supabase/supabase-js", "2.45.0", "Cliente oficial do Supabase para JavaScript"],
        ["dotenv", "16.4.0", "Carregamento de variáveis de ambiente (.env)"],
        ["node.js", "22.14.0", "Ambiente de execução JavaScript no servidor"],
    ],
    [4, 2.5, 9.5]
)

body("Para iniciar o servidor Express.js:")

code_block("cd api-express")
code_block("npm install")
code_block("npm start          # Servidor em http://localhost:3000")

section_h2("3.3", "Estrutura de Arquivos")

table(
    ["Arquivo", "Descrição"],
    [
        ["api-express/", "Diretório raiz da API Express"],
        ["  package.json", "Configuração do projeto e dependências npm"],
        ["  servidor.js", "Servidor principal com todas as rotas RESTful"],
        ["  supabase.js", "Cliente Supabase configurado com credenciais"],
        ["  .env", "Variáveis de ambiente (URL e chave do Supabase)"],
    ],
    [5, 11]
)

section_h2("3.4", "Arquitetura do Servidor")

body(
    "O servidor Express.js segue uma arquitetura monolítica simples, adequada "
    "para o MVP. O arquivo servidor.js concentra a configuração dos middlewares "
    "e todas as rotas organizadas por recurso. O arquivo supabase.js configura "
    "a conexão com o banco de dados Supabase, e cada endpoint realiza queries "
    "diretamente ao banco."
)

body("Os middlewares configurados são:")

bullet("cors(): permite que o front-end acesse a API de qualquer origem;")
bullet("express.json(): interpreta automaticamente o corpo das requisições em JSON.")

section_h2("3.5", "Endpoints Implementados")

section_h3("3.5.1", "Igrejas")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/igrejas", "Listar todas as igrejas"],
        ["GET", "/api/igrejas/:id", "Buscar igreja por ID"],
        ["POST", "/api/igrejas", "Criar nova igreja"],
        ["PUT", "/api/igrejas/:id", "Atualizar dados da igreja"],
        ["DELETE", "/api/igrejas/:id", "Remover igreja"],
    ],
    [2.5, 5, 8.5]
)

section_h3("3.5.2", "Membros")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/membros", "Listar membros (filtros: igrejaId, tipo)"],
        ["GET", "/api/membros/:id", "Buscar membro por ID"],
        ["POST", "/api/membros", "Criar novo membro"],
        ["PUT", "/api/membros/:id", "Atualizar dados do membro"],
        ["DELETE", "/api/membros/:id", "Remover membro"],
    ],
    [2.5, 5, 8.5]
)

section_h3("3.5.3", "Eventos")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/eventos", "Listar eventos (filtro: igrejaId)"],
        ["GET", "/api/eventos/:id", "Buscar evento por ID"],
        ["POST", "/api/eventos", "Criar novo evento"],
        ["PUT", "/api/eventos/:id", "Atualizar dados do evento"],
        ["DELETE", "/api/eventos/:id", "Remover evento"],
    ],
    [2.5, 5, 8.5]
)

section_h3("3.5.4", "Contribuições")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/contribuicoes", "Listar contribuições (filtros: igreja_id, membro_id, tipo)"],
        ["GET", "/api/contribuicoes/:id", "Buscar contribuição por ID"],
        ["POST", "/api/contribuicoes", "Registrar nova contribuição"],
        ["DELETE", "/api/contribuicoes/:id", "Remover contribuição"],
    ],
    [2.5, 5, 8.5]
)

section_h3("3.5.5", "Comunicados")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/comunicados", "Listar comunicados (filtro: igreja_id)"],
        ["GET", "/api/comunicados/:id", "Buscar comunicado por ID"],
        ["POST", "/api/comunicados", "Criar novo comunicado"],
        ["PUT", "/api/comunicados/:id", "Atualizar comunicado"],
        ["DELETE", "/api/comunicados/:id", "Remover comunicado"],
    ],
    [2.5, 5, 8.5]
)

section_h3("3.5.6", "Pedidos de Oração")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/pedidos-oracao", "Listar pedidos (filtros: igreja_id, membro_id, status)"],
        ["GET", "/api/pedidos-oracao/:id", "Buscar pedido por ID"],
        ["POST", "/api/pedidos-oracao", "Criar novo pedido de oração"],
        ["PUT", "/api/pedidos-oracao/:id", "Atualizar status do pedido"],
        ["DELETE", "/api/pedidos-oracao/:id", "Remover pedido"],
    ],
    [2.5, 5, 8.5]
)

section_h3("3.5.7", "Autenticação")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["POST", "/api/auth/registrar-igreja", "Cadastrar nova igreja (cria conta + igreja + pastor)"],
        ["POST", "/api/auth/registrar-membro", "Cadastrar novo membro (valida código da igreja)"],
        ["POST", "/api/auth/login", "Login unificado (igreja ou membro)"],
        ["POST", "/api/auth/recuperar-senha", "Enviar e-mail de recuperação de senha"],
    ],
    [2.5, 5, 8.5]
)

section_h2("3.6", "Exemplo de Requisição e Resposta")

body("Exemplo de criação de um membro via POST:", indent=False)

code_block('POST http://localhost:3000/api/membros')
code_block('Content-Type: application/json')
code_block('')
code_block('{')
code_block('  "nome_completo": "Ana Paula Silva",')
code_block('  "email": "ana@email.com",')
code_block('  "telefone": "(67) 99999-0004",')
code_block('  "tipo": "membro",')
code_block('  "igreja_id": "a1b2c3d4-e5f6-7890-abcd-ef1234567890"')
code_block('}')

blank(1)
body("Resposta (201 Created):", indent=False)

code_block('{')
code_block('  "id": "f7a8b9c0-d1e2-3456-abcd-ef1234567890",')
code_block('  "nome_completo": "Ana Paula Silva",')
code_block('  "email": "ana@email.com",')
code_block('  "telefone": "(67) 99999-0004",')
code_block('  "tipo": "membro",')
code_block('  "igreja_id": "a1b2c3d4-e5f6-7890-abcd-ef1234567890",')
code_block('  "criado_em": "2026-03-03T14:30:00.000Z",')
code_block('  "atualizado_em": "2026-03-03T14:30:00.000Z"')
code_block('}')

# ============================================================
# 4  FRAMEWORK 2 — FASTAPI
# ============================================================
page_break()
section_h1("4", "FRAMEWORK 2 \u2014 FASTAPI (PYTHON)")

section_h2("4.1", "Visão Geral")

body(
    "FastAPI é um framework moderno e de alta performance para criação de APIs "
    "em Python, lançado em 2018. Baseado no padrão ASGI (Asynchronous Server "
    "Gateway Interface) e no framework Starlette, utiliza tipagem nativa do Python "
    "e a biblioteca Pydantic para validação automática de dados. Seu principal "
    "diferencial é a geração automática de documentação interativa nos padrões "
    "Swagger (OpenAPI) e ReDoc."
)

section_h2("4.2", "Configuração e Dependências")

table(
    ["Pacote", "Versão", "Finalidade"],
    [
        ["fastapi", "0.115.0", "Framework web com validação automática e docs"],
        ["uvicorn", "0.30.0", "Servidor ASGI para executar a aplicação"],
        ["pydantic", "2.9.0", "Validação e serialização de dados com tipagem"],
        ["supabase", "2.9.1", "Cliente oficial do Supabase para Python"],
        ["python-dotenv", "1.0.1", "Carregamento de variáveis de ambiente (.env)"],
        ["python", "3.12.4", "Linguagem de programação"],
    ],
    [4, 2.5, 9.5]
)

body("Para iniciar o servidor FastAPI:")

code_block("cd api-fastapi")
code_block("pip install -r requirements.txt")
code_block("uvicorn servidor:app --reload    # Servidor em http://localhost:8000")
code_block("# Documentação automática em http://localhost:8000/docs")

section_h2("4.3", "Estrutura de Arquivos")

table(
    ["Arquivo", "Descrição"],
    [
        ["api-fastapi/", "Diretório raiz da API FastAPI"],
        ["  servidor.py", "Servidor principal com todas as rotas RESTful"],
        ["  modelos.py", "Schemas Pydantic para validação de dados"],
        ["  supabase_client.py", "Cliente Supabase configurado com credenciais"],
        ["  .env", "Variáveis de ambiente (URL e chave do Supabase)"],
        ["  requirements.txt", "Lista de dependências Python"],
    ],
    [5, 11]
)

section_h2("4.4", "Modelos de Dados (Pydantic)")

body(
    "O FastAPI utiliza modelos Pydantic para definir a estrutura dos dados "
    "esperados em cada requisição. Isso garante validação automática: se um "
    "campo obrigatório estiver ausente ou com tipo incorreto, a API retorna "
    "um erro detalhado automaticamente, sem necessidade de código manual."
)

body("Os modelos definidos para cada recurso são:")

table(
    ["Recurso", "Modelo de Criação", "Modelo de Atualização"],
    [
        ["Igreja", "IgrejaCriar", "IgrejaAtualizar"],
        ["Membro", "MembroCriar", "MembroAtualizar"],
        ["Evento", "EventoCriar", "EventoAtualizar"],
        ["Contribuição", "ContribuicaoCriar", "\u2014"],
        ["Comunicado", "ComunicadoCriar", "ComunicadoAtualizar"],
        ["Pedido de Oração", "PedidoOracaoCriar", "PedidoOracaoAtualizar"],
    ],
    [4, 5, 5]
)

section_h2("4.5", "Endpoints Implementados")

body(
    "Os endpoints do FastAPI são idênticos aos do Express.js em termos de "
    "funcionalidade, com a mesma estrutura de URLs e métodos HTTP. A diferença "
    "está na nomenclatura dos campos (snake_case no Python vs camelCase no "
    "JavaScript) e nos recursos adicionais como documentação automática e "
    "validação via Pydantic."
)

section_h3("4.5.1", "Igrejas")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/igrejas", "Listar todas as igrejas"],
        ["GET", "/api/igrejas/{igreja_id}", "Buscar igreja por ID"],
        ["POST", "/api/igrejas", "Criar nova igreja"],
        ["PUT", "/api/igrejas/{igreja_id}", "Atualizar dados da igreja"],
        ["DELETE", "/api/igrejas/{igreja_id}", "Remover igreja"],
    ],
    [2.5, 5.5, 8]
)

section_h3("4.5.2", "Membros")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/membros", "Listar membros (filtros: igreja_id, tipo)"],
        ["GET", "/api/membros/{membro_id}", "Buscar membro por ID"],
        ["POST", "/api/membros", "Criar novo membro"],
        ["PUT", "/api/membros/{membro_id}", "Atualizar dados do membro"],
        ["DELETE", "/api/membros/{membro_id}", "Remover membro"],
    ],
    [2.5, 5.5, 8]
)

section_h3("4.5.3", "Eventos")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/eventos", "Listar eventos (filtro: igreja_id)"],
        ["GET", "/api/eventos/{evento_id}", "Buscar evento por ID"],
        ["POST", "/api/eventos", "Criar novo evento"],
        ["PUT", "/api/eventos/{evento_id}", "Atualizar dados do evento"],
        ["DELETE", "/api/eventos/{evento_id}", "Remover evento"],
    ],
    [2.5, 5.5, 8]
)

section_h3("4.5.4", "Contribuições")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/contribuicoes", "Listar contribuições (filtros: igreja_id, membro_id, tipo)"],
        ["GET", "/api/contribuicoes/{id}", "Buscar contribuição por ID"],
        ["POST", "/api/contribuicoes", "Registrar nova contribuição"],
        ["DELETE", "/api/contribuicoes/{id}", "Remover contribuição"],
    ],
    [2.5, 5.5, 8]
)

section_h3("4.5.5", "Comunicados")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/comunicados", "Listar comunicados (filtro: igreja_id)"],
        ["GET", "/api/comunicados/{id}", "Buscar comunicado por ID"],
        ["POST", "/api/comunicados", "Criar novo comunicado"],
        ["PUT", "/api/comunicados/{id}", "Atualizar comunicado"],
        ["DELETE", "/api/comunicados/{id}", "Remover comunicado"],
    ],
    [2.5, 5.5, 8]
)

section_h3("4.5.6", "Pedidos de Oração")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["GET", "/api/pedidos-oracao", "Listar pedidos (filtros: igreja_id, membro_id, status)"],
        ["GET", "/api/pedidos-oracao/{id}", "Buscar pedido por ID"],
        ["POST", "/api/pedidos-oracao", "Criar novo pedido de oração"],
        ["PUT", "/api/pedidos-oracao/{id}", "Atualizar status do pedido"],
        ["DELETE", "/api/pedidos-oracao/{id}", "Remover pedido"],
    ],
    [2.5, 5.5, 8]
)

section_h3("4.5.7", "Autenticação")

table(
    ["Método", "Endpoint", "Descrição"],
    [
        ["POST", "/api/auth/registrar-igreja", "Cadastrar nova igreja (cria conta + igreja + pastor)"],
        ["POST", "/api/auth/registrar-membro", "Cadastrar novo membro (valida código da igreja)"],
        ["POST", "/api/auth/login", "Login unificado (igreja ou membro)"],
        ["POST", "/api/auth/recuperar-senha", "Enviar e-mail de recuperação de senha"],
    ],
    [2.5, 5.5, 8]
)

section_h2("4.6", "Documentação Automática (Swagger)")

body(
    "Um dos principais diferenciais do FastAPI é a geração automática de "
    "documentação interativa. Ao acessar http://localhost:8000/docs, o "
    "desenvolvedor tem acesso a uma interface Swagger completa onde é possível "
    "visualizar todos os endpoints, seus parâmetros, schemas de entrada e "
    "saída, e testar as requisições diretamente pelo navegador."
)

body(
    "A documentação também está disponível no formato ReDoc em "
    "http://localhost:8000/redoc, que apresenta uma visão mais detalhada "
    "e formatada dos endpoints e modelos de dados."
)

# ============================================================
# 5  COMPARATIVO ENTRE OS FRAMEWORKS
# ============================================================
page_break()
section_h1("5", "COMPARATIVO ENTRE OS FRAMEWORKS")

body(
    "A implementação das mesmas funcionalidades em dois frameworks distintos "
    "permite uma análise comparativa objetiva. A tabela a seguir resume as "
    "principais diferenças e semelhanças:"
)

table(
    ["Critério", "Express.js", "FastAPI"],
    [
        ["Linguagem", "JavaScript (Node.js)", "Python 3.12"],
        ["Paradigma", "Callbacks e middlewares", "Decorators e tipagem"],
        ["Validação de dados", "Manual (if/else)", "Automática (Pydantic)"],
        ["Documentação da API", "Manual (Swagger separado)", "Automática (Swagger/ReDoc)"],
        ["Performance", "Alta (event loop V8)", "Muito alta (ASGI assíncrono)"],
        ["Curva de aprendizado", "Baixa", "Baixa a média"],
        ["Ecossistema", "Muito amplo (npm)", "Amplo (pip)"],
        ["Integração com front-end", "Mesma linguagem (JS)", "Linguagem diferente"],
        ["Tratamento de erros", "Manual", "Automático + HTTPException"],
        ["Tipagem", "Opcional (sem TS)", "Nativa (type hints)"],
    ],
    [4, 6, 6]
)

section_h2("5.1", "Quando Usar Cada Framework")

body(
    "Express.js é mais indicado quando a equipe já trabalha com JavaScript "
    "no front-end e deseja manter uma única linguagem em todo o projeto "
    "(full-stack JavaScript). Sua simplicidade e vasto ecossistema de "
    "middlewares o tornam ideal para prototipagem rápida e MVPs."
)

body(
    "FastAPI é mais indicado quando se deseja validação automática de dados, "
    "documentação gerada automaticamente, e alta performance. A tipagem nativa "
    "do Python e os modelos Pydantic reduzem significativamente a quantidade "
    "de código de validação manual e tornam a API mais robusta e auto-documentada."
)

section_h2("5.2", "Escolha para o Projeto")

body(
    "Para o CongregaFiel, ambas as APIs foram implementadas como demonstração "
    "acadêmica de domínio em múltiplas tecnologias. Na prática, o Express.js "
    "será utilizado como API principal por manter a consistência com o "
    "ecossistema JavaScript já utilizado no front-end do projeto."
)

# ============================================================
# 6  BANCO DE DADOS — SUPABASE
# ============================================================
page_break()
section_h1("6", "BANCO DE DADOS \u2014 SUPABASE")

section_h2("6.1", "O que é o Supabase")

body(
    "O Supabase é uma plataforma de banco de dados na nuvem que funciona como "
    "uma alternativa ao Firebase do Google. Ele utiliza o PostgreSQL, um dos "
    "bancos de dados mais confiáveis e utilizados no mundo. Na prática, o "
    "Supabase oferece um lugar seguro na internet para guardar todas as "
    "informações do sistema \u2014 como os dados das igrejas, dos membros, dos "
    "eventos e das contribuições \u2014 sem precisar instalar nada no computador."
)

body(
    "Diferente da Sprint anterior, onde os dados eram salvos em arquivos de "
    "texto (JSON) dentro do próprio servidor, agora todas as informações ficam "
    "armazenadas em um banco de dados profissional na nuvem. Isso traz "
    "segurança, velocidade e a garantia de que os dados não serão perdidos."
)

section_h2("6.2", "Por que Supabase?")

bullet("Gratuito para projetos pequenos e médios;")
bullet("Banco de dados PostgreSQL profissional, usado por grandes empresas;")
bullet("Painel visual na internet para ver e gerenciar os dados;")
bullet("Funciona com qualquer linguagem de programação (JavaScript, Python, etc.);")
bullet("Segurança integrada com Row Level Security (controle de acesso por linha).")

section_h2("6.3", "Como o Supabase se conecta ao sistema")

body(
    "O site do Congrega Fiel (front-end) nunca acessa o banco de dados "
    "diretamente. Ele sempre faz pedidos para as APIs (Express.js ou FastAPI), "
    "que funcionam como intermediários. As APIs recebem o pedido, consultam o "
    "banco de dados no Supabase, e devolvem a resposta para o site. Isso "
    "garante que as credenciais de acesso ao banco fiquem protegidas no "
    "servidor, sem nunca serem expostas ao usuário."
)

body("O fluxo funciona assim:")

bullet("O usuário acessa o site e clica em algo (ex: ver lista de membros);")
bullet("O site envia um pedido para a API (Express.js ou FastAPI);")
bullet("A API recebe o pedido e consulta o Supabase;")
bullet("O Supabase retorna os dados do banco de dados;")
bullet("A API devolve os dados formatados para o site;")
bullet("O site exibe as informações na tela do usuário.")

section_h2("6.4", "Configuração de Segurança")

body(
    "Cada API possui um arquivo .env que contém duas informações essenciais: "
    "a URL do projeto Supabase e a chave secreta (service role key). Essa "
    "chave permite acesso total ao banco e por isso nunca é compartilhada "
    "publicamente. O arquivo .env está listado no .gitignore para garantir "
    "que não seja enviado ao repositório."
)

body(
    "Além disso, todas as tabelas possuem Row Level Security (RLS) habilitado, "
    "que é uma camada extra de proteção do PostgreSQL. Mesmo que alguém "
    "consiga a chave pública do projeto, não terá acesso aos dados sem as "
    "permissões corretas."
)

# ============================================================
# 7  MODELO DE DADOS
# ============================================================
page_break()
section_h1("7", "MODELO DE DADOS")

body(
    "O banco de dados no Supabase é composto por seis tabelas principais, "
    "todas com nomes em português. Cada tabela possui campos de auditoria "
    "(criado_em, atualizado_em, criado_por, atualizado_por) que registram "
    "automaticamente quando e por quem cada dado foi criado ou modificado. "
    "Os identificadores são UUIDs (códigos únicos universais) ao invés de "
    "números sequenciais, garantindo maior segurança."
)

section_h2("7.1", "Igrejas")

body(
    "Armazena os dados de cada igreja cadastrada no sistema, incluindo nome, "
    "endereço, código de identificação e informações do pastor responsável."
)

table(
    ["Campo", "Tipo", "Obrigatório", "Descrição"],
    [
        ["id", "UUID", "Auto", "Identificador único universal"],
        ["nome", "Texto", "Sim", "Nome da igreja"],
        ["endereco", "Texto", "Não", "Endereço completo"],
        ["descricao", "Texto", "Não", "Descrição da igreja"],
        ["codigo", "Texto", "Sim", "Código único (ex: CF1234)"],
        ["nome_pastor", "Texto", "Não", "Nome do pastor responsável"],
        ["email", "Texto", "Não", "E-mail da igreja"],
        ["senha_hash", "Texto", "Não", "Senha criptografada"],
        ["criado_em", "Data/Hora", "Auto", "Data de criação"],
        ["atualizado_em", "Data/Hora", "Auto", "Última atualização"],
    ],
    [3.5, 2.5, 2.5, 7.5]
)

section_h2("7.2", "Membros")

body(
    "Armazena os dados dos membros e pastores de cada igreja."
)

table(
    ["Campo", "Tipo", "Obrigatório", "Descrição"],
    [
        ["id", "UUID", "Auto", "Identificador único universal"],
        ["nome_completo", "Texto", "Sim", "Nome completo do membro"],
        ["email", "Texto", "Não", "E-mail do membro"],
        ["telefone", "Texto", "Não", "Telefone de contato"],
        ["tipo", "Texto", "Não", "pastor ou membro (padrão: membro)"],
        ["senha_hash", "Texto", "Não", "Senha criptografada"],
        ["igreja_id", "UUID", "Sim", "Igreja à qual pertence"],
        ["codigo_igreja", "Texto", "Não", "Código da igreja"],
        ["criado_em", "Data/Hora", "Auto", "Data de criação"],
        ["atualizado_em", "Data/Hora", "Auto", "Última atualização"],
    ],
    [3.5, 2.5, 2.5, 7.5]
)

section_h2("7.3", "Eventos")

body(
    "Armazena os eventos organizados por cada igreja."
)

table(
    ["Campo", "Tipo", "Obrigatório", "Descrição"],
    [
        ["id", "UUID", "Auto", "Identificador único universal"],
        ["titulo", "Texto", "Sim", "Título do evento"],
        ["descricao", "Texto", "Não", "Descrição do evento"],
        ["data", "Data", "Sim", "Data do evento"],
        ["horario", "Texto", "Não", "Horário do evento"],
        ["local", "Texto", "Não", "Local de realização"],
        ["igreja_id", "UUID", "Sim", "Igreja organizadora"],
        ["criado_em", "Data/Hora", "Auto", "Data de criação"],
        ["atualizado_em", "Data/Hora", "Auto", "Última atualização"],
    ],
    [3.5, 2.5, 2.5, 7.5]
)

section_h2("7.4", "Contribuições")

body(
    "Registra as contribuições financeiras dos membros (dízimos, ofertas, doações)."
)

table(
    ["Campo", "Tipo", "Obrigatório", "Descrição"],
    [
        ["id", "UUID", "Auto", "Identificador único universal"],
        ["membro_id", "UUID", "Sim", "Membro que contribuiu"],
        ["igreja_id", "UUID", "Sim", "Igreja que recebeu"],
        ["membro_nome", "Texto", "Não", "Nome do membro (referência rápida)"],
        ["tipo", "Texto", "Sim", "dízimo, oferta, doação ou outro"],
        ["valor", "Decimal", "Sim", "Valor da contribuição (> 0)"],
        ["data", "Data", "Não", "Data da contribuição"],
        ["descricao", "Texto", "Não", "Descrição adicional"],
        ["criado_em", "Data/Hora", "Auto", "Data de criação"],
        ["atualizado_em", "Data/Hora", "Auto", "Última atualização"],
    ],
    [3.5, 2.5, 2.5, 7.5]
)

section_h2("7.5", "Comunicados")

body(
    "Armazena os comunicados e avisos enviados pela igreja aos seus membros."
)

table(
    ["Campo", "Tipo", "Obrigatório", "Descrição"],
    [
        ["id", "UUID", "Auto", "Identificador único universal"],
        ["igreja_id", "UUID", "Sim", "Igreja que enviou o comunicado"],
        ["titulo", "Texto", "Sim", "Título do comunicado"],
        ["conteudo", "Texto", "Sim", "Conteúdo do comunicado"],
        ["prioridade", "Texto", "Não", "normal ou urgente"],
        ["criado_em", "Data/Hora", "Auto", "Data de criação"],
        ["atualizado_em", "Data/Hora", "Auto", "Última atualização"],
    ],
    [3.5, 2.5, 2.5, 7.5]
)

section_h2("7.6", "Pedidos de Oração")

body(
    "Armazena os pedidos de oração enviados pelos membros da igreja."
)

table(
    ["Campo", "Tipo", "Obrigatório", "Descrição"],
    [
        ["id", "UUID", "Auto", "Identificador único universal"],
        ["igreja_id", "UUID", "Sim", "Igreja do membro"],
        ["membro_id", "UUID", "Sim", "Membro que fez o pedido"],
        ["membro_nome", "Texto", "Não", "Nome do membro (referência rápida)"],
        ["pedido", "Texto", "Sim", "Texto do pedido de oração"],
        ["status", "Texto", "Não", "pendente, orado ou respondido"],
        ["criado_em", "Data/Hora", "Auto", "Data de criação"],
        ["atualizado_em", "Data/Hora", "Auto", "Última atualização"],
    ],
    [3.5, 2.5, 2.5, 7.5]
)

section_h2("7.7", "Diagrama de Relacionamentos")

body(
    "As tabelas se relacionam entre si da seguinte forma: cada membro pertence "
    "a uma igreja, cada evento é organizado por uma igreja, cada contribuição "
    "é feita por um membro para uma igreja, cada comunicado é enviado por uma "
    "igreja, e cada pedido de oração é feito por um membro dentro de uma igreja. "
    "Se uma igreja for removida, todos os dados relacionados (membros, eventos, "
    "contribuições, comunicados e pedidos) são removidos automaticamente."
)

# ============================================================
# 8  COMO FUNCIONA NA PRÁTICA
# ============================================================
page_break()
section_h1("8", "COMO FUNCIONA NA PRÁTICA")

body(
    "Para entender melhor como o sistema funciona, imagine o Congrega Fiel "
    "como um restaurante. O site que o usuário vê (front-end) é como o "
    "salão do restaurante \u2014 é onde o cliente faz seus pedidos. As APIs "
    "(Express.js e FastAPI) são como os garçons \u2014 eles recebem os pedidos "
    "e levam até a cozinha. O Supabase é a cozinha \u2014 é onde os dados são "
    "preparados e armazenados."
)

section_h2("8.1", "O que é um Framework de API?")

body(
    "Um framework de API é uma ferramenta que facilita a criação de um "
    "\"garçom digital\". Em vez de criar tudo do zero, o framework já "
    "vem com as peças prontas para receber pedidos da internet, processar "
    "esses pedidos e devolver respostas. No nosso projeto, temos dois "
    "garçons diferentes que fazem o mesmo trabalho, mas de formas distintas."
)

section_h2("8.2", "Express.js \u2014 O Garçom Rápido e Flexível")

body(
    "O Express.js é como um garçom experiente que fala a mesma língua "
    "do salão (JavaScript). Como o site do Congrega Fiel foi feito em "
    "JavaScript, usar o Express.js no servidor significa que toda a equipe "
    "trabalha com a mesma linguagem, facilitando a comunicação."
)

body("Exemplos práticos do que o Express.js faz no Congrega Fiel:")

bullet("Quando o pastor abre o painel e quer ver a lista de membros da sua "
       "igreja, o site envia um pedido para o Express.js. O Express recebe "
       "esse pedido, busca os membros no Supabase e devolve a lista formatada;")

bullet("Quando um membro faz uma contribuição (dízimo ou oferta), o site "
       "envia os dados para o Express.js, que salva o registro no Supabase "
       "com o valor, a data e o tipo de contribuição;")

bullet("Quando a igreja publica um comunicado urgente, o Express.js recebe "
       "o título e o conteúdo, salva no Supabase e, na próxima vez que um "
       "membro abrir o painel, o comunicado aparece na tela;")

bullet("Quando alguém faz um pedido de oração, o Express.js registra o "
       "pedido no Supabase com o status \"pendente\". O pastor pode depois "
       "atualizar o status para \"orado\" ou \"respondido\".")

section_h2("8.3", "FastAPI \u2014 O Garçom Organizado e Automático")

body(
    "O FastAPI é como um garçom que fala outra língua (Python), mas é "
    "extremamente organizado. Ele verifica automaticamente se o pedido "
    "está correto antes de enviar para a cozinha. Se alguém pedir um "
    "\"bolo\" no campo de \"valor\" (texto onde deveria ser número), o "
    "FastAPI recusa o pedido e explica exatamente o que está errado."
)

body("Exemplos práticos do que o FastAPI faz no Congrega Fiel:")

bullet("Quando o pastor cadastra um novo evento, o FastAPI verifica "
       "automaticamente se o título foi preenchido, se a data está no "
       "formato correto e se a igreja existe. Se algum campo estiver "
       "errado, retorna uma mensagem clara do que precisa ser corrigido;")

bullet("Quando alguém tenta registrar uma contribuição com valor zero "
       "ou negativo, o FastAPI bloqueia automaticamente e informa que o "
       "valor precisa ser maior que zero \u2014 sem precisar escrever código "
       "extra para isso;")

bullet("Quando o pastor quer ver apenas os comunicados urgentes, o "
       "FastAPI filtra automaticamente e retorna só os relevantes;")

bullet("O FastAPI gera automaticamente uma página de documentação "
       "(http://localhost:8000/docs) onde qualquer pessoa pode ver "
       "todos os serviços disponíveis e testá-los diretamente pelo "
       "navegador, sem precisar de ferramentas extras.")

section_h2("8.4", "Diferença entre Express.js e FastAPI na Prática")

body(
    "A principal diferença está em como cada um lida com os pedidos. "
    "O Express.js é mais simples e direto \u2014 ele confia que o pedido "
    "está correto e processa rapidamente. É como um garçom que anota "
    "o pedido e leva direto para a cozinha."
)

body(
    "O FastAPI é mais cuidadoso \u2014 ele verifica cada detalhe do pedido "
    "antes de processar. É como um garçom que confirma cada item, verifica "
    "se tem no cardápio e só então leva para a cozinha. Essa verificação "
    "extra acontece automaticamente, sem precisar programar manualmente."
)

body("Resumo das diferenças de forma simples:")

table(
    ["Característica", "Express.js", "FastAPI"],
    [
        ["Língua que usa", "JavaScript (mesma do site)", "Python (linguagem diferente)"],
        ["Velocidade de criar", "Muito rápido de montar", "Rápido, com bônus de organização"],
        ["Verificação de erros", "O programador precisa criar", "Já vem pronto automaticamente"],
        ["Documentação", "Precisa criar separadamente", "Gerada automaticamente"],
        ["Resultado final", "Funciona igual", "Funciona igual, com extras"],
    ],
    [4, 5.5, 6.5]
)

body(
    "No Congrega Fiel, ambos os frameworks fazem exatamente a mesma coisa: "
    "recebem pedidos do site, consultam o banco de dados Supabase e devolvem "
    "as informações. A diferença está no \"como\" fazem isso. É como dois "
    "caminhos diferentes que levam ao mesmo destino."
)

# ============================================================
# 9  AUTENTICAÇÃO COM SUPABASE
# ============================================================
page_break()
section_h1("9", "AUTENTICAÇÃO COM SUPABASE")

section_h2("9.1", "Visão Geral")

body(
    "A autenticação é o processo que garante que apenas pessoas autorizadas "
    "acessem o sistema. No Congrega Fiel, implementamos um sistema completo "
    "de autenticação utilizando o Supabase Auth, que funciona em conjunto "
    "com o banco de dados. Quando alguém cria uma conta ou faz login, os "
    "dados são salvos tanto no sistema de autenticação do Supabase quanto "
    "nas tabelas do banco de dados."
)

body(
    "O sistema possui dois tipos de usuário: igreja (pastor/administrador) "
    "e membro. Cada tipo tem acesso a diferentes funcionalidades do sistema. "
    "A igreja pode gerenciar membros, eventos, comunicados e contribuições. "
    "O membro pode visualizar informações, fazer contribuições e enviar "
    "pedidos de oração."
)

section_h2("9.2", "Como Funciona o Cadastro")

body(
    "O cadastro funciona de forma diferente para igrejas e membros. Quando "
    "um pastor cadastra sua igreja, o sistema cria automaticamente três "
    "coisas: uma conta de acesso (para fazer login), o registro da igreja "
    "no banco de dados (com um código único como CF1234), e o registro do "
    "pastor como primeiro membro da igreja. Tudo isso acontece em uma única "
    "ação, garantindo que os dados fiquem sempre sincronizados."
)

body(
    "Quando um membro quer se cadastrar, ele precisa informar o código da "
    "igreja que recebeu do pastor. O sistema verifica se esse código existe "
    "e, se estiver correto, cria a conta do membro já vinculada à igreja. "
    "Se o código estiver errado, o cadastro não é permitido, garantindo que "
    "apenas pessoas autorizadas se juntem a cada igreja."
)

body("O processo de cadastro segue este fluxo:")

bullet("O usuário preenche o formulário no site (nome, e-mail, senha);")
bullet("O site envia os dados para a API (Express.js ou FastAPI);")
bullet("A API cria a conta no Supabase Auth (sistema de autenticação);")
bullet("A API salva os dados complementares nas tabelas do banco de dados;")
bullet("Se houver qualquer erro, tudo é desfeito automaticamente (rollback);")
bullet("O usuário recebe uma confirmação e é redirecionado para o login.")

section_h2("9.3", "Como Funciona o Login")

body(
    "O login é unificado — tanto igrejas quanto membros usam a mesma tela "
    "de login, informando apenas e-mail e senha. O sistema identifica "
    "automaticamente o tipo de usuário e redireciona para o painel correto: "
    "o painel administrativo (para igrejas) ou o painel do membro."
)

body("O processo de login segue este fluxo:")

bullet("O usuário informa e-mail e senha na tela de login;")
bullet("O site envia os dados para a API;")
bullet("A API verifica as credenciais no Supabase Auth;")
bullet("Se estiver correto, a API busca os dados completos do usuário no banco;")
bullet("A API retorna os dados do usuário e um token de acesso (access token);")
bullet("O site salva a sessão localmente e sincroniza os dados da igreja;")
bullet("O usuário é redirecionado para o painel correto.")

body(
    "O token de acesso é como um crachá digital que o usuário recebe ao "
    "fazer login. Esse crachá é enviado automaticamente em todas as próximas "
    "requisições, permitindo que a API saiba quem está fazendo cada pedido "
    "sem precisar pedir a senha novamente."
)

section_h2("9.4", "Recuperação de Senha")

body(
    "Se o usuário esquecer a senha, ele pode solicitar uma recuperação. "
    "O sistema envia um e-mail automático (via Supabase) com um link para "
    "criar uma nova senha. Por segurança, o site sempre mostra a mensagem "
    "\"E-mail enviado com sucesso\", mesmo que o e-mail não esteja cadastrado "
    "— isso evita que alguém descubra quais e-mails existem no sistema."
)

section_h2("9.5", "Como o Frontend se Conecta à API")

body(
    "O site do Congrega Fiel possui um serviço dedicado chamado ApiServico "
    "que gerencia toda a comunicação com a API. Esse serviço é responsável "
    "por enviar os dados de cadastro, login e recuperação de senha, além de "
    "sincronizar automaticamente os dados da igreja após o login."
)

body(
    "Quando o usuário faz login com sucesso, o sistema carrega automaticamente "
    "todos os dados da igreja para o navegador: membros, eventos, contribuições, "
    "comunicados e pedidos de oração. Isso permite que as telas do painel "
    "funcionem rapidamente, pois os dados já estão disponíveis localmente. "
    "Quando o usuário adiciona ou remove algo, a alteração é salva tanto "
    "localmente quanto enviada para a API em segundo plano."
)

section_h2("9.6", "Segurança da Autenticação")

body(
    "O sistema de autenticação implementa diversas camadas de segurança:"
)

bullet("As senhas nunca são armazenadas diretamente — o Supabase Auth utiliza "
       "criptografia bcrypt para proteger cada senha;")
bullet("As credenciais do banco de dados (chave secreta) ficam apenas no "
       "servidor, nunca são expostas ao navegador do usuário;")
bullet("O token de acesso tem validade limitada e é renovado automaticamente;")
bullet("Se o cadastro falhar após a criação da conta, o sistema desfaz tudo "
       "automaticamente (rollback) para não deixar dados inconsistentes;")
bullet("A recuperação de senha não revela se um e-mail está cadastrado ou não;")
bullet("Cada requisição à API inclui o token de autenticação no cabeçalho.")

table(
    ["Recurso de Segurança", "Descrição"],
    [
        ["Supabase Auth", "Sistema profissional de autenticação com criptografia"],
        ["bcrypt", "Algoritmo de hash para proteger senhas armazenadas"],
        ["Access Token (JWT)", "Token temporário que identifica o usuário logado"],
        ["Rollback automático", "Desfaz operações incompletas para manter consistência"],
        ["Row Level Security", "Proteção por linha no banco de dados"],
        ["Variáveis de ambiente", "Credenciais protegidas em arquivo .env no servidor"],
    ],
    [5, 11]
)

# ============================================================
# 10  INSTRUÇÕES DE EXECUÇÃO
# ============================================================
page_break()
section_h1("10", "INSTRUÇÕES DE EXECUÇÃO")

section_h2("10.1", "Pré-requisitos")

bullet("Node.js versão 18 ou superior (para Express.js);")
bullet("Python versão 3.10 ou superior (para FastAPI);")
bullet("npm (gerenciador de pacotes do Node.js);")
bullet("pip (gerenciador de pacotes do Python);")
bullet("Conta no Supabase com projeto criado (https://supabase.com).")

section_h2("10.2", "Configuração do Banco de Dados")

body("Antes de iniciar qualquer API, é necessário configurar o Supabase:")

bullet("Acesse o painel do Supabase e abra o SQL Editor;")
bullet("Execute o arquivo database/schema.sql para criar as tabelas;")
bullet("Copie a URL e a Service Role Key do projeto (Settings > API);")
bullet("Crie um arquivo .env em cada pasta de API com as credenciais.")

blank(1)
body("Exemplo do arquivo .env:", indent=False)

code_block("SUPABASE_URL=https://seu-projeto.supabase.co")
code_block("SUPABASE_SECRET_KEY=sua-service-role-key-aqui")

section_h2("10.3", "Executando a API Express.js")

code_block("# 1. Acessar o diretório da API")
code_block("cd api-express")
blank(1)
code_block("# 2. Instalar dependências")
code_block("npm install")
blank(1)
code_block("# 3. Iniciar o servidor")
code_block("npm start")
blank(1)
code_block("# Servidor disponível em http://localhost:3000")
code_block("# Testar: curl http://localhost:3000/api/igrejas")

section_h2("10.4", "Executando a API FastAPI")

code_block("# 1. Acessar o diretório da API")
code_block("cd api-fastapi")
blank(1)
code_block("# 2. Instalar dependências")
code_block("pip install -r requirements.txt")
blank(1)
code_block("# 3. Iniciar o servidor")
code_block("uvicorn servidor:app --reload")
blank(1)
code_block("# Servidor disponível em http://localhost:8000")
code_block("# Documentação Swagger: http://localhost:8000/docs")

# ============================================================
# 11  CONSIDERAÇÕES FINAIS
# ============================================================
page_break()
section_h1("11", "CONSIDERAÇÕES FINAIS")

body(
    "A Sprint 2 representou um avanço significativo na infraestrutura do "
    "projeto Congrega Fiel. A implementação de duas Web APIs REST demonstra "
    "a versatilidade da equipe no uso de diferentes linguagens e frameworks. "
    "O Express.js provou ser uma solução eficiente e direta para a criação de "
    "APIs em JavaScript, enquanto o FastAPI trouxe benefícios significativos "
    "em termos de validação automática e documentação integrada."
)

body(
    "A migração para o Supabase como banco de dados representou uma evolução "
    "importante: os dados deixaram de ser armazenados em arquivos locais (JSON) "
    "e passaram a ficar em um banco de dados profissional na nuvem. Isso traz "
    "maior segurança, confiabilidade e permite que múltiplos servidores acessem "
    "os mesmos dados simultaneamente."
)

body(
    "A implementação da autenticação com Supabase Auth completou o ciclo de "
    "segurança do sistema. Agora, tanto igrejas quanto membros possuem contas "
    "reais com senhas criptografadas, tokens de acesso e recuperação de senha "
    "por e-mail. O cadastro de igrejas gera automaticamente um código único "
    "que permite que membros se vinculem à igreja correta, enquanto o login "
    "unificado identifica automaticamente o tipo de usuário."
)

body(
    "Ambas as implementações seguem o padrão RESTful com os métodos HTTP "
    "corretos (GET, POST, PUT, DELETE) e oferecem os mesmos recursos "
    "funcionais: gestão de igrejas, membros, eventos, contribuições "
    "financeiras, comunicados e pedidos de oração, além de autenticação "
    "completa. Todos os módulos do sistema estão cobertos por endpoints "
    "completos em ambas as APIs, integrados ao frontend do projeto."
)

body(
    "A integração entre frontend e backend foi implementada através do "
    "ApiServico, um módulo JavaScript que gerencia toda a comunicação com "
    "a API. O sistema sincroniza automaticamente os dados do Supabase para "
    "o navegador após o login, e envia as alterações feitas pelo usuário "
    "de volta para o servidor em segundo plano, garantindo uma experiência "
    "fluida e dados sempre atualizados."
)

body(
    "Como próximos passos para a Sprint 3, está previsto o refinamento da "
    "interface do painel administrativo e do painel dos membros, testes de "
    "integração completos e possíveis melhorias na experiência do usuário."
)

# ============================================================
# REFERÊNCIAS
# ============================================================
page_break()
section_h1("", "REFERÊNCIAS")

for ref in [
    'EXPRESS.JS. Express \u2014 Node.js web application framework. '
    'Disponível em: https://expressjs.com/. Acesso em: 03 mar. 2026.',

    'FASTAPI. FastAPI \u2014 modern, fast web framework for building APIs with Python. '
    'Disponível em: https://fastapi.tiangolo.com/. Acesso em: 03 mar. 2026.',

    'SUPABASE. Supabase \u2014 The Open Source Firebase Alternative. '
    'Disponível em: https://supabase.com/docs. Acesso em: 03 mar. 2026.',

    'POSTGRESQL. PostgreSQL \u2014 The World\'s Most Advanced Open Source Database. '
    'Disponível em: https://www.postgresql.org/. Acesso em: 03 mar. 2026.',

    'MOZILLA DEVELOPER NETWORK. HTTP Methods. '
    'Disponível em: https://developer.mozilla.org/pt-BR/docs/Web/HTTP/Methods. '
    'Acesso em: 03 mar. 2026.',

    'NODE.JS. Node.js \u2014 JavaScript runtime. '
    'Disponível em: https://nodejs.org/. Acesso em: 03 mar. 2026.',

    'PYDANTIC. Pydantic \u2014 Data validation using Python type annotations. '
    'Disponível em: https://docs.pydantic.dev/. Acesso em: 03 mar. 2026.',

    'FIELDING, R. T. Architectural Styles and the Design of Network-based '
    'Software Architectures. Dissertação (Doutorado) \u2014 University of California, '
    'Irvine, 2000.',

    'PRESSMAN, R. S.; MAXIM, B. R. Engenharia de software: uma abordagem '
    'profissional. 9. ed. Porto Alegre: AMGH, 2021.',

    'SOMMERVILLE, I. Engenharia de software. 10. ed. São Paulo: Pearson, 2019.',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(ref)
    r.font.size = Pt(SZ); r.font.name = FONT_NAME
    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

# ============================================================
# SALVAR
# ============================================================
out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs")
os.makedirs(out_dir, exist_ok=True)
out = os.path.join(out_dir, "Semana 2 - Web API Express.js e FastAPI.docx")
doc.save(out)
print(f"OK  -> {out}")
print("Ao abrir no Word, pressione Ctrl+A e F9 para gerar o sumário linkado.")
