# -*- coding: utf-8 -*-
"""
Documentação ABNT — Congrega Fiel (MVP)
Faculdade INSTED · Campo Grande/MS · Design comercial
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import os

# ============================================================
# PALETA COMERCIAL
# ============================================================
CLR_PRIMARY   = "1B2A4A"   # azul marinho profundo
CLR_ACCENT    = "C5A55A"   # dourado elegante
CLR_LIGHT_BG  = "F0EDE6"   # creme claro (zebra rows)
CLR_WHITE     = "FFFFFF"
CLR_DARK_TEXT = "1B2A4A"
CLR_MID_GRAY  = "6B7280"
CLR_LIGHT_LN  = "D4C99E"   # dourado claro (linhas)
CLR_TABLE_HDR = "1B2A4A"   # header tabela

# ============================================================
# CONFIG GERAL
# ============================================================
FONT_NAME = "Times New Roman"
SZ = 12          # corpo
SZ_TITLE = 14
SZ_COVER = 16
SZ_BIG = 28
LSPACING = 1.5

INST  = "FACULDADE INSTED"
CURSO = "Curso Superior de Tecnologia em Análise e Desenvolvimento de Sistemas"
CITY  = "Campo Grande \u2013 MS"
YEAR  = "2026"

doc = Document()

# Margens ABNT
for sec in doc.sections:
    sec.top_margin    = Cm(3)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2)
    sec.page_height   = Cm(29.7)
    sec.page_width    = Cm(21)

# Auto-update fields (sumário)
uf = OxmlElement("w:updateFields")
uf.set(qn("w:val"), "true")
doc.settings.element.append(uf)

# ============================================================
# ESTILOS BASE
# ============================================================
sn = doc.styles["Normal"]
sn.font.name = FONT_NAME
sn.font.size = Pt(SZ)
sn.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
sn.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
sn.paragraph_format.line_spacing = LSPACING
sn.paragraph_format.space_after  = Pt(0)
sn.paragraph_format.space_before = Pt(0)
rPr = sn.element.get_or_add_rPr()
rPr.append(parse_xml(
    f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" '
    f'w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>'
))

# Heading styles (para TOC)
for lvl, name in enumerate(["Heading 1", "Heading 2", "Heading 3"], 1):
    hs = doc.styles[name]
    hs.font.name = FONT_NAME
    hs.font.size = Pt(SZ)
    hs.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    hs.font.bold = lvl <= 2
    hs.font.italic = lvl == 3
    hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hs.paragraph_format.line_spacing = LSPACING
    hs.paragraph_format.space_before = Pt(18 if lvl == 1 else 12 if lvl == 2 else 8)
    hs.paragraph_format.space_after  = Pt(6 if lvl == 1 else 4)
    hs.paragraph_format.first_line_indent = Cm(0)
    hs.paragraph_format.left_indent = Cm(0)
    pPr = hs.element.get_or_add_pPr()
    num = pPr.find(qn("w:numPr"))
    if num is not None:
        pPr.remove(num)
    hrPr = hs.element.get_or_add_rPr()
    hrPr.append(parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" '
        f'w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>'
    ))
    hrPr.append(parse_xml(f'<w:color {nsdecls("w")} w:val="{CLR_PRIMARY}"/>'))

# ============================================================
# HELPERS
# ============================================================
def blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(); r.font.size = Pt(SZ); r.font.name = FONT_NAME

def centered(text, size=SZ, bold=False, upper=False, color=CLR_DARK_TEXT):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LSPACING
    r = p.add_run(text.upper() if upper else text)
    r.font.size = Pt(size); r.font.name = FONT_NAME; r.bold = bold
    r.font.color.rgb = RGBColor(int(color[:2],16), int(color[2:4],16), int(color[4:],16))
    return p

def body(text, size=SZ, bold=False, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LSPACING
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = FONT_NAME; r.bold = bold
    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return p

def colored_line(color=CLR_ACCENT, thickness=12):
    """Linha horizontal colorida decorativa."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{thickness}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def section_h1(number, title):
    """Heading 1 com linha dourada embaixo."""
    text = f"{number}  {title.upper()}" if number else title.upper()
    p = doc.add_paragraph(text, style="Heading 1")
    for r in p.runs:
        r.font.size = Pt(SZ); r.font.name = FONT_NAME; r.bold = True
        r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    # Borda inferior dourada no parágrafo
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="2" w:color="{CLR_ACCENT}"/>'
        f'</w:pBdr>'
    ))
    return p

def section_h2(number, title):
    p = doc.add_paragraph(f"{number}  {title}", style="Heading 2")
    for r in p.runs:
        r.font.size = Pt(SZ); r.font.name = FONT_NAME; r.bold = True
        r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return p

def section_h3(number, title):
    p = doc.add_paragraph(f"{number}  {title}", style="Heading 3")
    for r in p.runs:
        r.font.size = Pt(SZ); r.font.name = FONT_NAME
        r.bold = False; r.italic = True
        r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = LSPACING
    ind = 1.25 + (level * 0.75)
    p.paragraph_format.left_indent = Cm(ind)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    sym = "\u2022" if level == 0 else "\u25E6"
    r = p.add_run(f"{sym}  {text}")
    r.font.size = Pt(SZ); r.font.name = FONT_NAME
    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"

    # Cabeçalho — fundo azul marinho, texto dourado
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after  = Pt(4)
        rr = pp.add_run(h)
        rr.font.size = Pt(10); rr.font.name = FONT_NAME; rr.bold = True
        rr.font.color.rgb = RGBColor(0xC5, 0xA5, 0x5A)  # dourado
        c._element.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CLR_TABLE_HDR}" w:val="clear"/>')
        )

    # Dados — zebra creme
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ""
            pp = c.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pp.paragraph_format.space_before = Pt(3)
            pp.paragraph_format.space_after  = Pt(3)
            rr = pp.add_run(str(val))
            rr.font.size = Pt(10); rr.font.name = FONT_NAME
            rr.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
            if ri % 2 == 0:
                c._element.get_or_add_tcPr().append(
                    parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CLR_WHITE}" w:val="clear"/>')
                )
            else:
                c._element.get_or_add_tcPr().append(
                    parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CLR_LIGHT_BG}" w:val="clear"/>')
                )

    # Bordas da tabela — dourado claro
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'  <w:left   w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'  <w:right  w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{CLR_LIGHT_LN}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    blank(1)
    return t

def toc_field():
    """Sumário automático linkado do Word."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = LSPACING

    r1 = p.add_run()
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    r1._element.append(fc1)

    r2 = p.add_run()
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
    ins.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._element.append(ins)

    r3 = p.add_run()
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate")
    r3._element.append(fc2)

    r4 = p.add_run("Abra no Word e pressione F9 para gerar o sumário linkado.")
    r4.font.size = Pt(SZ); r4.font.name = FONT_NAME
    r4.font.color.rgb = RGBColor(0x6B, 0x72, 0x80); r4.italic = True

    r5 = p.add_run()
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end")
    r5._element.append(fc3)

def page_break():
    doc.add_page_break()

def cover_band(color, height_cm=1.2):
    """Faixa colorida horizontal (simulada com tabela de 1 célula)."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.text = ""
    cell.height = Cm(height_cm)
    # Fundo
    cell._element.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    )
    # Sem bordas
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left   w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right  w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    ))
    # Largura total
    tblPr.append(parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    ))

# ============================================================
#                       CAPA COMERCIAL
# ============================================================
blank(1)
cover_band(CLR_PRIMARY, 0.6)
blank(1)
centered(INST, size=SZ_COVER, bold=True, upper=True, color=CLR_PRIMARY)
colored_line(CLR_ACCENT, 6)
centered(CURSO, size=11, bold=False, color=CLR_MID_GRAY)
blank(4)
centered("CONGREGA FIEL", size=SZ_BIG, bold=True, upper=True, color=CLR_PRIMARY)
blank(1)
# Linha dourada grossa
colored_line(CLR_ACCENT, 18)
blank(1)
centered("Sistema Web para Gestão de", size=SZ_TITLE, bold=False, color=CLR_MID_GRAY)
centered("Comunidades Eclesiásticas", size=SZ_TITLE, bold=False, color=CLR_MID_GRAY)
blank(3)

# Nomes com estilo
for nome in [
    "Catieli Gama Cora",
    "Fernando Alves da Nóbrega",
    "Gabriel Franklin Barcellos",
    "Jhenniffer Lopes da Silva Vargas",
    "João Pedro Aranda",
]:
    centered(nome, size=SZ, bold=False, color=CLR_PRIMARY)

blank(4)
cover_band(CLR_ACCENT, 0.15)
blank(1)
centered(CITY, size=SZ, bold=False, color=CLR_MID_GRAY)
centered(YEAR, size=SZ, bold=True, color=CLR_PRIMARY)

# ============================================================
#                    FOLHA DE ROSTO
# ============================================================
page_break()
blank(2)
centered("CONGREGA FIEL", size=22, bold=True, upper=True, color=CLR_PRIMARY)
colored_line(CLR_ACCENT, 10)
centered("Sistema Web para Gestão de Comunidades Eclesiásticas",
         size=SZ_TITLE, bold=False, color=CLR_MID_GRAY)
blank(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0
r = p.add_run(
    f"Documentação técnica do Produto Mínimo Viável (MVP) apresentada como "
    f"requisito parcial para aprovação no {CURSO} da {INST}."
)
r.font.size = Pt(10); r.font.name = FONT_NAME
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

blank(3)
centered("Equipe de Desenvolvimento", size=SZ, bold=True, color=CLR_PRIMARY)
colored_line(CLR_ACCENT, 4)
blank(1)
for nome in [
    "Catieli Gama Cora",
    "Fernando Alves da Nóbrega",
    "Gabriel Franklin Barcellos",
    "Jhenniffer Lopes da Silva Vargas",
    "João Pedro Aranda",
]:
    centered(nome, size=SZ, color=CLR_PRIMARY)

blank(5)
cover_band(CLR_ACCENT, 0.1)
blank(1)
centered(CITY, size=SZ, color=CLR_MID_GRAY)
centered(YEAR, size=SZ, bold=True, color=CLR_PRIMARY)

# ============================================================
#                       SUMÁRIO
# ============================================================
page_break()
centered("SUMÁRIO", size=SZ_TITLE, bold=True, upper=True, color=CLR_PRIMARY)
colored_line(CLR_ACCENT, 8)
blank(1)
toc_field()

# ============================================================
# 1  INTRODUÇÃO
# ============================================================
page_break()
section_h1("1", "INTRODUÇÃO")

body(
    "O presente documento descreve a documentação técnica inicial do projeto "
    "Congrega Fiel, um sistema web desenvolvido como Produto Mínimo Viável (MVP) "
    "com o objetivo de oferecer uma solução tecnológica para a gestão de comunidades "
    f"eclesiásticas. O projeto foi concebido no âmbito do {CURSO} "
    f"da {INST} e visa atender a uma demanda real identificada no "
    "cotidiano de igrejas e comunidades religiosas."
)

body(
    "A gestão de uma comunidade religiosa envolve uma série de atividades "
    "administrativas e pastorais que, em muitos casos, ainda são realizadas de forma "
    "manual ou descentralizada. O cadastro de membros, o controle de contribuições "
    "financeiras, a organização de eventos e a comunicação entre líderes e fiéis são "
    "processos que demandam tempo, organização e, frequentemente, geram retrabalho "
    "ou perda de informações."
)

body(
    "Nesse contexto, o Congrega Fiel surge como uma plataforma digital que centraliza "
    "essas atividades em um único ambiente web, proporcionando praticidade, segurança "
    "e transparência na administração eclesiástica. O sistema funciona como uma espécie "
    "de rede social privada para cada igreja, onde pastores podem cadastrar suas "
    "congregações, gerenciar seus membros e acompanhar as atividades da comunidade, "
    "enquanto os fiéis têm acesso a informações relevantes sobre eventos, contribuições "
    "e comunicados."
)

body(
    "Este documento apresenta o escopo do MVP, os requisitos funcionais e não "
    "funcionais, as tecnologias selecionadas, a arquitetura proposta, a definição "
    "de papéis da equipe e o cronograma de desenvolvimento, servindo como base "
    "para todo o ciclo de vida do projeto."
)

# ============================================================
# 2  PROBLEMA IDENTIFICADO
# ============================================================
page_break()
section_h1("2", "PROBLEMA IDENTIFICADO")

body(
    "As comunidades religiosas, especialmente igrejas de pequeno e médio porte, "
    "enfrentam desafios significativos na gestão de suas atividades administrativas "
    "e pastorais. Foram identificados os seguintes problemas recorrentes:"
)

for prob in [
    "Gestão manual de membros: o cadastro e o acompanhamento dos fiéis são realizados "
    "em planilhas, cadernos ou sistemas genéricos, resultando em dados desatualizados, "
    "duplicados ou perdidos.",

    "Controle financeiro precário: as contribuições (dízimos, ofertas e doações) "
    "são registradas sem padronização, dificultando a prestação de contas.",

    "Comunicação fragmentada: a divulgação de eventos e comunicados depende de "
    "redes sociais genéricas, gerando ruído e perda de informações.",

    "Ausência de centralização: não existe uma plataforma unificada para gerenciar "
    "membros, finanças, eventos e comunicação em um único lugar.",

    "Falta de privacidade: o uso de redes sociais públicas expõe informações "
    "sensíveis da comunidade.",
]:
    bullet(prob)

body(
    "Diante desse cenário, evidencia-se a necessidade de uma ferramenta digital "
    "específica que atenda às particularidades da gestão eclesiástica."
)

# ============================================================
# 3  SOLUÇÃO PROPOSTA
# ============================================================
page_break()
section_h1("3", "SOLUÇÃO PROPOSTA")

body(
    "Propõe-se o desenvolvimento do Congrega Fiel, um sistema web que funciona "
    "como uma plataforma privada de gestão para comunidades eclesiásticas, "
    "projetada com foco na simplicidade e acessibilidade."
)

section_h2("3.1", "Conceito do Sistema")

body(
    "O Congrega Fiel opera como uma rede social privada para igrejas. "
    "Cada congregação possui seu espaço digital gerenciado pelo pastor, "
    "com dois perfis de acesso:"
)

bullet(
    "Pastor/Líder (Administrador): acesso completo \u2014 cadastrar igreja, "
    "gerenciar membros, registrar pagamentos, criar eventos e comunicados."
)
bullet(
    "Fiel/Membro: acesso ao seu perfil, eventos, histórico de contribuições "
    "e comunicados da igreja."
)

section_h2("3.2", "Principais Funcionalidades do MVP")

for f in [
    "Cadastro e autenticação de usuários (pastores e fiéis);",
    "Cadastro e configuração de igrejas por parte do pastor;",
    "Cadastro e gestão de membros vinculados a cada igreja;",
    "Registro e controle de contribuições financeiras;",
    "Criação e divulgação de eventos da comunidade;",
    "Painel administrativo para o pastor;",
    "Área do fiel com acesso às suas informações.",
]:
    bullet(f)

section_h2("3.3", "Diferenciais")

body(
    "Privacidade por padrão (cada igreja em ambiente isolado), interface "
    "intuitiva para diferentes perfis de usuário, e funcionalidades voltadas "
    "especificamente para o contexto eclesiástico."
)

# ============================================================
# 4  ESCOPO DO PROJETO (MVP)
# ============================================================
page_break()
section_h1("4", "ESCOPO DO PROJETO (MVP)")

body(
    "O escopo está delimitado a um Produto Mínimo Viável contendo as "
    "funcionalidades essenciais para validação do conceito."
)

section_h2("4.1", "Funcionalidades Incluídas")

section_h3("4.1.1", "Módulo de Autenticação e Cadastro")
for i in ["Cadastro de pastor com validação de dados;",
          "Cadastro de fiel vinculado a uma igreja;",
          "Login e logout com controle de sessão;",
          "Recuperação de senha;",
          "Diferenciação de perfis (pastor e fiel)."]:
    bullet(i)

section_h3("4.1.2", "Módulo de Gestão da Igreja")
for i in ["Cadastro de nova igreja pelo pastor;",
          "Edição de informações (nome, endereço, descrição);",
          "Painel da igreja com dados resumidos."]:
    bullet(i)

section_h3("4.1.3", "Módulo de Gestão de Membros")
for i in ["Cadastro de novos membros pelo pastor;",
          "Listagem com filtros de busca;",
          "Edição e exclusão de cadastros;",
          "Visualização de perfil individual."]:
    bullet(i)

section_h3("4.1.4", "Módulo Financeiro")
for i in ["Registro de contribuições (dízimos, ofertas, doações);",
          "Histórico de pagamentos por membro;",
          "Relatório financeiro resumido para o pastor;",
          "Consulta pelo fiel (apenas as próprias contribuições)."]:
    bullet(i)

section_h3("4.1.5", "Módulo de Eventos")
for i in ["Criação de eventos pelo pastor;",
          "Listagem de eventos para os fiéis;",
          "Detalhes do evento (data, horário, local, descrição)."]:
    bullet(i)

section_h2("4.2", "Fora do Escopo do MVP")

for i in ["Aplicativo mobile nativo;",
          "Chat em tempo real;",
          "Integração com gateways de pagamento;",
          "Relatórios avançados em PDF;",
          "Notificações push;",
          "Módulo de escola bíblica;",
          "Integração com calendários externos."]:
    bullet(i)

# ============================================================
# 5  REQUISITOS DO SISTEMA
# ============================================================
page_break()
section_h1("5", "REQUISITOS DO SISTEMA")

section_h2("5.1", "Requisitos Funcionais")

table(
    ["ID", "Requisito", "Descrição"],
    [
        ["RF01", "Cadastro de Pastor", "Cadastro com nome, e-mail, telefone e senha."],
        ["RF02", "Cadastro de Igreja", "Pastor cadastra igreja com nome, endereço e descrição."],
        ["RF03", "Cadastro de Fiel", "Pastor cadastra fiéis vinculados à igreja."],
        ["RF04", "Autenticação", "Login/logout com validação de credenciais."],
        ["RF05", "Recuperação de Senha", "Recuperação via e-mail cadastrado."],
        ["RF06", "Painel do Pastor", "Dashboard com resumo de membros, finanças e eventos."],
        ["RF07", "Gestão de Membros", "Listar, buscar, editar e excluir membros."],
        ["RF08", "Contribuições", "Registrar contribuições vinculadas a membros."],
        ["RF09", "Histórico Financeiro", "Histórico filtrável por período e tipo."],
        ["RF10", "Criação de Eventos", "Criar eventos com título, data, horário e local."],
        ["RF11", "Listagem de Eventos", "Listar eventos ordenados por data."],
        ["RF12", "Área do Fiel", "Acesso a perfil, contribuições e eventos."],
        ["RF13", "Controle de Acesso", "Funcionalidades admin restritas ao pastor."],
    ],
    [2, 3.5, 10.5]
)

section_h2("5.2", "Requisitos Não Funcionais")

table(
    ["ID", "Categoria", "Descrição"],
    [
        ["RNF01", "Usabilidade", "Interface intuitiva para diferentes perfis de usuários."],
        ["RNF02", "Responsividade", "Adaptável a desktop, tablet e smartphone."],
        ["RNF03", "Desempenho", "Carregamento em até 3 segundos."],
        ["RNF04", "Segurança", "Senhas criptografadas; prevenção contra XSS e SQL Injection."],
        ["RNF05", "Compatibilidade", "Chrome, Firefox, Edge e Safari (versões recentes)."],
        ["RNF06", "Manutenibilidade", "Separação clara entre HTML, CSS e JavaScript."],
        ["RNF07", "Disponibilidade", "99% de uptime em produção."],
    ],
    [2, 3, 11]
)

# ============================================================
# 6  TECNOLOGIAS UTILIZADAS
# ============================================================
page_break()
section_h1("6", "TECNOLOGIAS UTILIZADAS")

section_h2("6.1", "Front-End")

for tech, desc in [
    ("HTML5", "Linguagem de marcação para estruturação semântica das páginas."),
    ("CSS3", "Estilização com Flexbox e Grid Layout para responsividade."),
    ("JavaScript (ES6+)", "Lógica de interação, validações, requisições assíncronas e manipulação do DOM."),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = LSPACING
    p.paragraph_format.first_line_indent = Cm(1.25)
    rb = p.add_run(f"{tech}: ")
    rb.font.size = Pt(SZ); rb.font.name = FONT_NAME; rb.bold = True
    rb.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
    rd = p.add_run(desc)
    rd.font.size = Pt(SZ); rd.font.name = FONT_NAME
    rd.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

section_h2("6.2", "Back-End e Armazenamento")

body(
    "Para o MVP, a persistência utiliza LocalStorage/SessionStorage do navegador. "
    "A simulação da API é feita com JSON Server, com migração futura prevista "
    "para Express com banco de dados."
)

section_h2("6.3", "Ferramentas de Apoio")

table(
    ["Ferramenta", "Finalidade"],
    [
        ["Visual Studio Code", "Editor de código principal"],
        ["Git / GitHub", "Controle de versão e repositório remoto"],
        ["Figma", "Prototipação e design de interface"],
        ["Trello", "Gerenciamento de tarefas"],
    ],
    [5, 11]
)

# ============================================================
# 7  ARQUITETURA ORIENTADA A SERVIÇOS
# ============================================================
page_break()
section_h1("7", "ARQUITETURA ORIENTADA A SERVIÇOS")

body(
    "A Arquitetura Orientada a Serviços (SOA) é um modelo no qual as "
    "funcionalidades são disponibilizadas como serviços independentes que "
    "se comunicam por protocolos padronizados. No Congrega Fiel, o front-end "
    "consome dados de uma camada de serviços via Web API."
)

section_h2("7.1", "Web Services: SOAP e REST")

body("Existem dois modelos principais de Web Services:")

bullet(
    "SOAP (Simple Object Access Protocol): protocolo baseado em XML com "
    "formato rígido de mensagens (WSDL). Indicado para cenários que exigem "
    "alta segurança e transações complexas."
)
bullet(
    "REST (Representational State Transfer): estilo arquitetural leve que "
    "utiliza os métodos HTTP e formatos como JSON. Amplamente adotado em "
    "aplicações web e mobile modernas."
)

body(
    "O Congrega Fiel adota REST por ser mais adequado a aplicações web, "
    "mais simples de implementar e por utilizar JSON, que possui integração "
    "nativa com JavaScript."
)

section_h2("7.2", "Protocolo HTTP e Métodos RESTful")

body(
    "O HTTP é o protocolo de comunicação da web. No padrão RESTful, cada "
    "recurso é acessado por uma URL e manipulado pelos métodos:"
)

table(
    ["Método", "Ação", "Exemplo"],
    [
        ["GET", "Consultar/listar", "GET /api/membros"],
        ["POST", "Criar recurso", "POST /api/membros"],
        ["PUT", "Atualizar recurso", "PUT /api/membros/1"],
        ["DELETE", "Remover recurso", "DELETE /api/membros/1"],
    ],
    [2.5, 4, 9.5]
)

section_h2("7.3", "Web API e Frameworks")

body(
    "Uma Web API expõe funcionalidades via HTTP. Dois frameworks foram "
    "estudados para o projeto:"
)

bullet(
    "Node.js com Express: framework minimalista para criar rotas RESTful "
    "com baixa curva de aprendizado. Vasto ecossistema via npm."
)
bullet(
    "JSON Server: simula uma API REST completa a partir de um arquivo JSON, "
    "ideal para desenvolvimento e testes do MVP."
)

body(
    "Na fase inicial utiliza-se JSON Server; a migração para Express "
    "com banco de dados ocorrerá em versões futuras."
)

section_h2("7.4", "Estrutura de Diretórios")

table(
    ["Caminho", "Descrição"],
    [
        ["congregafiel/", "Diretório raiz"],
        ["  index.html", "Landing page"],
        ["  pages/", "Páginas HTML"],
        ["  css/", "Folhas de estilo"],
        ["  js/", "Scripts JavaScript"],
        ["  assets/", "Imagens e ícones"],
        ["  db.json", "Banco simulado (JSON Server)"],
    ],
    [5.5, 10.5]
)

# ============================================================
# 8  EQUIPE E PAPÉIS
# ============================================================
page_break()
section_h1("8", "DEFINIÇÃO DA EQUIPE E PAPÉIS")

body(
    "A equipe é composta por cinco integrantes em três frentes: "
    "Documentação, Front-End e Back-End."
)

section_h2("8.1", "Composição")

table(
    ["Integrante", "Área(s)", "Responsabilidades"],
    [
        ["Catieli Gama Cora", "Documentação", "Elaboração e revisão da documentação"],
        ["Fernando Alves da Nóbrega", "Documentação / Front-End", "Documentação e interfaces"],
        ["Gabriel Franklin Barcellos", "Front-End / Back-End", "Interfaces e lógica de negócios"],
        ["Jhenniffer Lopes da Silva Vargas", "Documentação / Front-End", "Documentação e interfaces"],
        ["João Pedro Aranda", "Back-End", "Lógica de negócios e persistência"],
    ],
    [5, 4, 7]
)

section_h2("8.2", "Frentes de Trabalho")

section_h3("8.2.1", "Documentação")
body("Fernando, Jhenniffer e Catieli — requisitos, escopo, diagramas e relatórios.")

section_h3("8.2.2", "Front-End")
body("Fernando, Jhenniffer e Gabriel — interfaces, responsividade e acessibilidade.")

section_h3("8.2.3", "Back-End")
body("João e Gabriel — lógica de negócios, validações e persistência de dados.")

# ============================================================
# 9  CRONOGRAMA
# ============================================================
page_break()
section_h1("9", "CRONOGRAMA DO MVP")

body(
    "O desenvolvimento está organizado em quatro sprints semanais com início "
    "em 24 de fevereiro de 2026."
)

table(
    ["Sprint", "Período", "Atividades"],
    [
        ["Sprint 1", "24/02 \u2013 02/03", "Documentação, requisitos, escopo, prototipação (Figma)"],
        ["Sprint 2", "03/03 \u2013 09/03", "HTML das páginas, CSS responsivo, módulo de autenticação"],
        ["Sprint 3", "10/03 \u2013 16/03", "Módulos de membros, financeiro e eventos; integração com API"],
        ["Sprint 4", "17/03 \u2013 23/03", "Integração final, testes, documentação e apresentação"],
    ],
    [2.5, 3, 10.5]
)

# ============================================================
# 10  CONSIDERAÇÕES FINAIS
# ============================================================
page_break()
section_h1("10", "CONSIDERAÇÕES FINAIS")

body(
    "O Congrega Fiel aplica os conhecimentos do curso superior em uma solução "
    "real para a gestão de comunidades religiosas. O MVP contempla cadastro de "
    "igrejas e membros, controle financeiro, eventos e diferenciação de perfis."
)

body(
    "A utilização de HTML, CSS e JavaScript garante compatibilidade, baixo custo "
    "e facilidade de manutenção. Como próximos passos: back-end com banco de "
    "dados, integração com pagamento online e aplicativo mobile."
)

# ============================================================
# REFERÊNCIAS
# ============================================================
page_break()
section_h1("", "REFERÊNCIAS")

for ref in [
    'ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. ABNT NBR 6023: informação e '
    'documentação \u2014 referências \u2014 elaboração. Rio de Janeiro: ABNT, 2018.',

    'ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. ABNT NBR 14724: informação e '
    'documentação \u2014 trabalhos acadêmicos \u2014 apresentação. Rio de Janeiro: ABNT, 2011.',

    'MOZILLA DEVELOPER NETWORK. HTML: Linguagem de Marcação de HiperTexto. '
    'Disponível em: https://developer.mozilla.org/pt-BR/docs/Web/HTML. '
    'Acesso em: 24 fev. 2026.',

    'MOZILLA DEVELOPER NETWORK. CSS: Folhas de Estilo em Cascata. '
    'Disponível em: https://developer.mozilla.org/pt-BR/docs/Web/CSS. '
    'Acesso em: 24 fev. 2026.',

    'MOZILLA DEVELOPER NETWORK. JavaScript. '
    'Disponível em: https://developer.mozilla.org/pt-BR/docs/Web/JavaScript. '
    'Acesso em: 24 fev. 2026.',

    'PRESSMAN, R. S.; MAXIM, B. R. Engenharia de software: uma abordagem '
    'profissional. 9. ed. Porto Alegre: AMGH, 2021.',

    'RIES, E. A startup enxuta. São Paulo: Leya, 2012.',

    'SOMMERVILLE, I. Engenharia de software. 10. ed. São Paulo: Pearson, 2019.',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(ref)
    r.font.size = Pt(SZ); r.font.name = FONT_NAME
    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

# ============================================================
# SALVAR
# ============================================================
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "Documentacao_CongregaFiel_MVP.docx")
doc.save(out)
print(f"OK  -> {out}")
print("Ao abrir no Word, pressione Ctrl+A e F9 para gerar o sumário.")
