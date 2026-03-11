from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ==================== ESTILOS GLOBAIS ====================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Estilo dos headings
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)

# Margens ABNT
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)


# ==================== FUNCOES AUXILIARES ====================
def add_cover_line(text, size=12, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Times New Roman'
    return p


def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run.bold = True
    return h


def add_body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def make_table(headers, data):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return table


def add_toc():
    """Insere campo de sumario automatico (TOC) com links."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = p.add_run('(Clique com botao direito e selecione "Atualizar campo", ou pressione F9)')
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(128, 128, 128)

    run5 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)


# ==================== CAPA ====================
for _ in range(5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('FACULDADE INSTED', 14, True, 4)
add_cover_line('Curso Superior de Tecnologia em Analise e Desenvolvimento de Sistemas', 12, False, 0)

for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('CONGREGA FIEL', 16, True, 4)
add_cover_line('Sistema Web para Gestao de Comunidades Eclesiasticas', 13, False, 0)

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('Relatorio da Sprint 2 - Semana 2', 13, True, 0)

for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('Catieli Gama Cora', 12, False, 2)
add_cover_line('Fernando Alves da Nobrega', 12, False, 2)
add_cover_line('Gabriel Franklin Barcellos', 12, False, 2)
add_cover_line('Jhenniffer Lopes da Silva Vargas', 12, False, 2)
add_cover_line('Joao Pedro Aranda', 12, False, 0)

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('Campo Grande - MS', 12, False, 2)
add_cover_line('2026', 12, False, 0)
doc.add_page_break()

# ==================== FOLHA DE ROSTO ====================
for _ in range(5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('CONGREGA FIEL', 16, True, 4)
add_cover_line('Sistema Web para Gestao de Comunidades Eclesiasticas', 13, False, 0)

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(0)
run = p.add_run('Relatorio tecnico da Sprint 2 apresentado como requisito parcial para aprovacao no Curso Superior de Tecnologia em Analise e Desenvolvimento de Sistemas da FACULDADE INSTED.')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('Equipe de Desenvolvimento', 12, True, 6)
add_cover_line('Catieli Gama Cora', 12, False, 2)
add_cover_line('Fernando Alves da Nobrega', 12, False, 2)
add_cover_line('Gabriel Franklin Barcellos', 12, False, 2)
add_cover_line('Jhenniffer Lopes da Silva Vargas', 12, False, 2)
add_cover_line('Joao Pedro Aranda', 12, False, 0)

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('Campo Grande - MS', 12, False, 2)
add_cover_line('2026', 12, False, 0)
doc.add_page_break()

# ==================== SUMARIO ====================
add_heading_custom('SUMARIO', 1)
add_toc()
doc.add_page_break()

# ==================== 1 INTRODUCAO ====================
add_heading_custom('1  INTRODUCAO', 1)
add_body('Este documento apresenta o relatorio da Sprint 2 do projeto Congrega Fiel, correspondente ao periodo de 03 a 09 de marco de 2026. Nesta etapa, o sistema saiu do ambiente local e passou a funcionar online, com quatro grandes entregas:', indent=True)
add_bullet('Hospedagem do site na internet com Firebase Hosting;')
add_bullet('Criacao do banco de dados na nuvem com Supabase (PostgreSQL);')
add_bullet('Construcao de uma API com Express.js (JavaScript);')
add_bullet('Construcao de uma API equivalente com FastAPI (Python).')
add_body('Na Sprint 1, o sistema funcionava apenas no computador do desenvolvedor. Agora, qualquer pessoa com acesso a internet pode utilizar o sistema, e os dados ficam salvos de forma permanente em um servidor na nuvem.', indent=True)

# ==================== 2 OBJETIVOS ====================
add_heading_custom('2  OBJETIVOS DA SPRINT 2', 1)
add_body('Os objetivos definidos para esta sprint foram:', indent=True)
add_bullet('Colocar o site no ar, acessivel pela internet;')
add_bullet('Criar um banco de dados real para guardar informacoes de forma permanente;')
add_bullet('Criar duas APIs que conectam o site ao banco de dados;')
add_bullet('Implementar login e cadastro com seguranca;')
add_bullet('Organizar o codigo do front-end com servicos compartilhados.')

# ==================== 3 FIREBASE HOSTING ====================
add_heading_custom('3  FIREBASE HOSTING', 1)
add_body('O Firebase Hosting e um servico gratuito do Google que permite publicar sites na internet. Funciona como um endereco online para o sistema. Os arquivos do site (HTML, CSS, JavaScript) sao enviados para os servidores do Google atraves de um comando no terminal (firebase deploy), e o Google distribui esses arquivos em servidores ao redor do mundo, garantindo carregamento rapido.', indent=True)
add_body('O servico fornece automaticamente um certificado de seguranca HTTPS, representado pelo cadeado no navegador, que indica conexao segura. Apos o deploy, o sistema ficou acessivel no endereco https://congregafiel.web.app.', indent=True)

# ==================== 4 SUPABASE ====================
add_heading_custom('4  SUPABASE - BANCO DE DADOS', 1)
add_body('O Supabase e uma plataforma que oferece banco de dados PostgreSQL na nuvem. Funciona como um armario digital onde o sistema guarda todas as informacoes de forma permanente. Antes, os dados ficavam salvos apenas no navegador do usuario e eram perdidos ao trocar de computador. Agora, as informacoes ficam no servidor e podem ser acessadas de qualquer dispositivo.', indent=True)

add_heading_custom('4.1  Estrutura do Banco', 2)
add_body('O banco foi organizado em 6 tabelas, cada uma responsavel por um tipo de informacao:', indent=True)

make_table(
    ['Tabela', 'O que Guarda', 'Exemplo'],
    [
        ['igrejas', 'Dados das igrejas cadastradas', 'Nome, endereco, codigo, pastor'],
        ['membros', 'Pessoas vinculadas a uma igreja', 'Nome, e-mail, telefone, tipo'],
        ['eventos', 'Atividades e programacoes', 'Culto de Domingo, 10h, Templo Principal'],
        ['contribuicoes', 'Dizimos, ofertas e doacoes', 'Maria - Dizimo - R$ 500,00'],
        ['comunicados', 'Avisos e informacoes', 'Retiro Espiritual dias 20-22 de marco'],
        ['pedidos_oracao', 'Pedidos de oracao dos membros', 'Oracao pela saude da minha mae'],
    ]
)

add_heading_custom('4.2  Relacionamentos e Seguranca', 2)
add_body('As tabelas se conectam entre si: cada membro pertence a uma igreja, cada evento pertence a uma igreja, e cada contribuicao pertence a um membro e a uma igreja. Ao excluir uma igreja, todos os dados associados sao removidos automaticamente.', indent=True)
add_body('O banco possui Row Level Security (RLS), mecanismo que controla quem pode ver e modificar cada registro. As credenciais de acesso ficam em arquivos protegidos (.env) que nao sao compartilhados publicamente.', indent=True)

# ==================== 5 EXPRESS.JS ====================
add_heading_custom('5  API COM EXPRESS.JS', 1)
add_body('Uma API (Interface de Programacao de Aplicacoes) funciona como um intermediario entre o site e o banco de dados. Quando o usuario clica em "Ver Eventos", o site faz um pedido a API, que busca os dados no banco e devolve para exibir na tela. A primeira API foi construida com Express.js, o framework web mais popular para JavaScript.', indent=True)

add_heading_custom('5.1  Funcionamento', 2)
add_body('A API recebe pedidos HTTP e responde com dados em formato JSON. Por exemplo: ao acessar a rota /api/eventos, a API consulta o banco de dados e retorna a lista de eventos com titulo, data, horario e local de cada um. A API possui 33 rotas organizadas por recurso:', indent=True)

make_table(
    ['Recurso', 'Operacoes', 'Rotas'],
    [
        ['Igrejas', 'Listar, buscar, criar, atualizar, remover', '5'],
        ['Membros', 'Listar, buscar, criar, atualizar, remover', '5'],
        ['Eventos', 'Listar, buscar, criar, atualizar, remover', '5'],
        ['Contribuicoes', 'Listar, buscar, criar, remover', '4'],
        ['Comunicados', 'Listar, buscar, criar, atualizar, remover', '5'],
        ['Pedidos de Oracao', 'Listar, buscar, criar, atualizar, remover', '5'],
        ['Autenticacao', 'Cadastro igreja, cadastro membro, login, recuperar senha', '4'],
    ]
)

# ==================== 6 FASTAPI ====================
add_heading_custom('6  API COM FASTAPI', 1)
add_body('A segunda API foi construida com FastAPI, um framework moderno para Python. Ela possui as mesmas 33 rotas e acessa o mesmo banco de dados, mas oferece duas vantagens adicionais:', indent=True)
add_bullet('Validacao automatica: se alguem enviar dados incorretos (campo obrigatorio vazio, formato errado), a API recusa e explica o erro automaticamente, sem precisar programar essa verificacao;')
add_bullet('Documentacao automatica: ao acessar /docs no navegador, aparece uma pagina interativa (Swagger) onde e possivel visualizar todas as rotas e testa-las diretamente.')
add_body('A biblioteca Pydantic define modelos que especificam exatamente quais dados sao esperados em cada requisicao. Por exemplo, para criar um evento, o modelo exige titulo, data e identificador da igreja. Se qualquer campo estiver faltando, a API retorna uma mensagem de erro clara automaticamente.', indent=True)

# ==================== 7 COMPARATIVO ====================
add_heading_custom('7  COMPARATIVO ENTRE OS FRAMEWORKS', 1)
add_body('A implementacao da mesma API em dois frameworks diferentes permite uma analise comparativa:', indent=True)

make_table(
    ['Criterio', 'Express.js', 'FastAPI'],
    [
        ['Linguagem', 'JavaScript', 'Python'],
        ['Validacao de dados', 'Manual', 'Automatica (Pydantic)'],
        ['Documentacao da API', 'Manual', 'Automatica (Swagger + ReDoc)'],
        ['Curva de aprendizado', 'Baixa', 'Baixa a media'],
        ['Vantagem principal', 'Mesma linguagem do front-end', 'Validacao e docs automaticas'],
    ]
)

add_body('O Express.js e ideal quando a equipe ja usa JavaScript no front-end e deseja manter uma linguagem unica. O FastAPI e ideal quando se deseja validacao automatica e documentacao sem esforco adicional. Ambas acessam o mesmo banco de dados e retornam os mesmos dados, demonstrando a versatilidade da equipe.', indent=True)

# ==================== 8 AUTENTICACAO ====================
add_heading_custom('8  AUTENTICACAO', 1)
add_body('Autenticacao e o processo que garante que apenas pessoas autorizadas acessem o sistema. O Congrega Fiel utiliza o Supabase Auth, um servico profissional que cuida de toda a seguranca: criptografia de senhas, tokens de acesso e recuperacao de senha por e-mail.', indent=True)

add_heading_custom('8.1  Cadastro', 2)
add_body('O sistema possui dois tipos de cadastro. Para igrejas, o pastor informa nome, nome da igreja, e-mail e senha. O sistema cria a conta, gera um codigo unico para a igreja (exemplo: CF1234) e registra o pastor como administrador. Para membros, a pessoa informa nome, e-mail, codigo da igreja e senha. O sistema verifica se o codigo existe e vincula o membro a igreja correta.', indent=True)

add_heading_custom('8.2  Login e Recuperacao de Senha', 2)
add_body('O login e unificado: pastores e membros usam a mesma tela. O sistema identifica automaticamente o tipo de usuario e redireciona para o painel correto. Apos o login, e gerado um token de acesso (como um cracha digital temporario) enviado automaticamente em todas as requisicoes seguintes.', indent=True)
add_body('Para recuperacao de senha, o usuario informa seu e-mail e recebe um link para criar nova senha. Por seguranca, a mensagem de confirmacao aparece sempre, mesmo que o e-mail nao exista no sistema.', indent=True)

# ==================== 9 ARQUITETURA FRONT-END ====================
add_heading_custom('9  ARQUITETURA DO FRONT-END', 1)
add_body('O front-end foi organizado com tres servicos compartilhados, modulos de codigo que centralizam funcoes usadas por todas as paginas, evitando repeticao:', indent=True)

make_table(
    ['Servico', 'Responsabilidade', 'Exemplo de Uso'],
    [
        ['Sessao', 'Controlar quem esta logado', 'Verificar permissao de acesso a pagina'],
        ['Interface', 'Funcoes visuais reutilizaveis', 'Exibir notificacoes, formatar datas e valores'],
        ['API', 'Comunicacao com o servidor', 'Buscar eventos, enviar pedido de oracao'],
    ]
)

add_body('Com essa organizacao, cada pagina nao precisa reescrever codigo de autenticacao, formatacao ou comunicacao com o servidor. Basta chamar o servico correspondente.', indent=True)

# ==================== 10 DEPLOY ONLINE ====================
add_heading_custom('10  DEPLOY ONLINE', 1)
add_body('Para o sistema funcionar permanentemente sem depender de um computador local ligado, todos os componentes foram publicados online:', indent=True)

make_table(
    ['Componente', 'Plataforma', 'Endereco'],
    [
        ['Site (front-end)', 'Firebase Hosting', 'https://congregafiel.web.app'],
        ['API Express.js', 'Vercel', 'https://api-express-tau.vercel.app'],
        ['API FastAPI', 'Vercel', 'https://api-fastapi.vercel.app'],
        ['Banco de dados', 'Supabase', 'PostgreSQL na nuvem'],
    ]
)

add_body('O front-end detecta automaticamente o ambiente: em desenvolvimento local, conecta ao localhost; em producao, conecta a API hospedada no Vercel.', indent=True)

# ==================== 11 ESTRUTURA DO PROJETO ====================
add_heading_custom('11  ESTRUTURA DO PROJETO', 1)
add_body('A estrutura de diretorios do projeto ficou organizada da seguinte forma:', indent=True)

make_table(
    ['Pasta', 'Descricao'],
    [
        ['public/', 'Site completo (HTML, CSS, JavaScript)'],
        ['public/autenticacao/', 'Login, cadastro e recuperacao de senha'],
        ['public/igreja/', 'Painel administrativo do pastor (6 paginas)'],
        ['public/membros/', 'Painel dos membros (6 paginas)'],
        ['public/js/servicos/', 'Servicos compartilhados (sessao, interface, API)'],
        ['api-express/', 'API REST com Express.js + Supabase'],
        ['api-fastapi/', 'API REST com FastAPI + Supabase'],
        ['database/', 'Schema SQL do banco de dados'],
    ]
)

# ==================== 12 CONSIDERACOES FINAIS ====================
add_heading_custom('12  CONSIDERACOES FINAIS', 1)
add_body('A Sprint 2 marcou a transicao do projeto da fase de prototipacao para um sistema funcional completo. O site saiu do computador local e passou a ser acessivel pela internet. Os dados deixaram de ser temporarios e passaram a ser permanentes, salvos em um banco de dados na nuvem.', indent=True)
add_body('Foram construidas duas APIs REST completas, uma em JavaScript (Express.js) e outra em Python (FastAPI), ambas conectadas ao mesmo banco de dados Supabase. Isso demonstra a capacidade da equipe em trabalhar com diferentes linguagens e frameworks para resolver o mesmo problema.', indent=True)
add_body('O sistema de autenticacao garante que apenas usuarios autorizados acessem o sistema, com senhas criptografadas, tokens de acesso e recuperacao de senha por e-mail. A organizacao do codigo com servicos compartilhados no front-end eliminou a duplicacao e criou uma base solida para evolucoes futuras.', indent=True)

doc.add_page_break()

# ==================== REFERENCIAS ====================
add_heading_custom('REFERENCIAS', 1)
refs = [
    'EXPRESS.JS. Express - Node.js web application framework. Disponivel em: https://expressjs.com/. Acesso em: 03 mar. 2026.',
    'FASTAPI. FastAPI - modern, fast web framework for building APIs with Python. Disponivel em: https://fastapi.tiangolo.com/. Acesso em: 03 mar. 2026.',
    'FIREBASE. Firebase Hosting Documentation. Disponivel em: https://firebase.google.com/docs/hosting. Acesso em: 03 mar. 2026.',
    'SUPABASE. Supabase - The Open Source Firebase Alternative. Disponivel em: https://supabase.com/docs. Acesso em: 03 mar. 2026.',
    'POSTGRESQL. PostgreSQL Documentation. Disponivel em: https://www.postgresql.org/docs/. Acesso em: 03 mar. 2026.',
    'PYDANTIC. Pydantic - Data validation using Python type annotations. Disponivel em: https://docs.pydantic.dev/. Acesso em: 03 mar. 2026.',
    'VERCEL. Vercel - Develop. Preview. Ship. Disponivel em: https://vercel.com/docs. Acesso em: 03 mar. 2026.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

output_path = os.path.join('docs', 'Semana 2.docx')
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
