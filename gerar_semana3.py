from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ==================== ESTILOS GLOBAIS ====================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Estilo dos headings
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)

# Margens ABNT
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)


# ==================== FUNCOES AUXILIARES ====================
def add_cover_line(text, size=12, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Times New Roman'
    return p


def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run.bold = True
    return h


def add_body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def make_table(headers, data):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return table


def add_toc():
    """Insere campo de sumario automatico (TOC) com links."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = p.add_run('(Clique com botao direito e selecione "Atualizar campo", ou pressione F9)')
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(128, 128, 128)

    run5 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)


# ==================== CAPA ====================
for _ in range(5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('FACULDADE INSTED', 14, True, 4)
add_cover_line('Curso Superior de Tecnologia em Analise e Desenvolvimento de Sistemas', 12, False, 0)

for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('CONGREGA FIEL', 16, True, 4)
add_cover_line('Sistema Web para Gestao de Comunidades Eclesiasticas', 13, False, 0)

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('Relatorio da Sprint 3 - Semana 3', 13, True, 0)

for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('Catieli Gama Cora', 12, False, 2)
add_cover_line('Fernando Alves da Nobrega', 12, False, 2)
add_cover_line('Gabriel Franklin Barcellos', 12, False, 2)
add_cover_line('Jhenniffer Lopes da Silva Vargas', 12, False, 2)
add_cover_line('Joao Pedro Aranda', 12, False, 0)

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('Campo Grande - MS', 12, False, 2)
add_cover_line('2026', 12, False, 0)
doc.add_page_break()

# ==================== FOLHA DE ROSTO ====================
for _ in range(5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('CONGREGA FIEL', 16, True, 4)
add_cover_line('Sistema Web para Gestao de Comunidades Eclesiasticas', 13, False, 0)

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(0)
run = p.add_run('Relatorio tecnico da Sprint 3 apresentado como requisito parcial para aprovacao no Curso Superior de Tecnologia em Analise e Desenvolvimento de Sistemas da FACULDADE INSTED.')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('Equipe de Desenvolvimento', 12, True, 6)
add_cover_line('Catieli Gama Cora', 12, False, 2)
add_cover_line('Fernando Alves da Nobrega', 12, False, 2)
add_cover_line('Gabriel Franklin Barcellos', 12, False, 2)
add_cover_line('Jhenniffer Lopes da Silva Vargas', 12, False, 2)
add_cover_line('Joao Pedro Aranda', 12, False, 0)

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

add_cover_line('Campo Grande - MS', 12, False, 2)
add_cover_line('2026', 12, False, 0)
doc.add_page_break()

# ==================== SUMARIO ====================
add_heading_custom('SUMARIO', 1)
add_toc()
doc.add_page_break()

# ==================== 1 INTRODUCAO ====================
add_heading_custom('1  INTRODUCAO', 1)
add_body('Este documento apresenta o relatorio da Sprint 3 do projeto Congrega Fiel, correspondente ao periodo de 10 a 16 de marco de 2026. Nesta etapa, o sistema recebeu melhorias significativas na experiencia do usuario e integracoes com servicos web externos, com cinco grandes entregas:', indent=True)
add_bullet('Integracao com Web Service de mapas (OpenStreetMap) via biblioteca Leaflet.js;')
add_bullet('Geracao de QR Code para convite de membros no painel da igreja;')
add_bullet('Criacao da pagina Linha do Tempo com visualizacao cronologica de eventos;')
add_bullet('Categorizacao de eventos por tipo (culto, estudo, conferencia, especial);')
add_bullet('Melhorias gerais na interface, nos menus de navegacao e nos filtros de pedidos de oracao.')
add_body('Na Sprint 2, o sistema foi publicado online com banco de dados na nuvem e duas APIs REST. Agora, na Sprint 3, o foco foi enriquecer a experiencia do usuario com recursos visuais interativos e consumo de APIs externas gratuitas, alem de aprimorar funcionalidades existentes.', indent=True)

# ==================== 2 OBJETIVOS ====================
add_heading_custom('2  OBJETIVOS DA SPRINT 3', 1)
add_body('Os objetivos definidos para esta sprint foram:', indent=True)
add_bullet('Integrar um Web Service gratuito de mapas para facilitar o cadastro de membros;')
add_bullet('Permitir que igrejas compartilhem convites via QR Code;')
add_bullet('Criar uma visualizacao cronologica (linha do tempo) dos eventos da igreja;')
add_bullet('Adicionar categorizacao de eventos por tipo;')
add_bullet('Melhorar a navegacao e os filtros nas paginas existentes;')
add_bullet('Criar pagina de perfil para membros.')

# ==================== 3 WEB SERVICE DE MAPAS ====================
add_heading_custom('3  WEB SERVICE DE MAPAS', 1)
add_body('Um Web Service e um servico disponibilizado na internet que permite a comunicacao entre sistemas diferentes. No Congrega Fiel, foi integrado o OpenStreetMap, um servico gratuito e colaborativo de mapas, por meio da biblioteca JavaScript Leaflet.js. Essa integracao constitui o consumo de uma API externa (Web Service REST), atendendo ao requisito de utilizacao de servicos web no projeto.', indent=True)

add_heading_custom('3.1  Funcionamento da Integracao', 2)
add_body('O OpenStreetMap fornece tiles (blocos de imagem) de mapas atraves de requisicoes HTTP. A biblioteca Leaflet.js consome esses tiles e renderiza um mapa interativo diretamente no navegador. A cada movimentacao do mapa (zoom, arrasto), novas requisicoes sao feitas ao servidor do OpenStreetMap para carregar os blocos de imagem correspondentes a area visivel.', indent=True)
add_body('Alem dos tiles de mapa, o sistema utiliza a API de Geolocalizacao do navegador (Geolocation API) para obter a posicao geografica do usuario com sua permissao. Essas coordenadas sao usadas para calcular a distancia ate cada igreja cadastrada, facilitando a escolha da mais proxima.', indent=True)

add_heading_custom('3.2  Mapa no Cadastro de Membros', 2)
add_body('Na pagina de cadastro de membros, o mapa exibe marcadores para cada igreja cadastrada no sistema. As igrejas sao carregadas a partir do endpoint publico GET /api/igrejas/publicas, que retorna nome, endereco, codigo, latitude e longitude de cada igreja sem exigir autenticacao.', indent=True)
add_body('O layout da pagina foi reorganizado em duas colunas: o mapa interativo a esquerda e o formulario de cadastro a direita. O usuario pode localizar sua igreja de tres formas:', indent=True)
add_bullet('Navegando pelo mapa e clicando no marcador da igreja desejada;')
add_bullet('Utilizando o campo de busca para filtrar igrejas por nome, endereco ou codigo;')
add_bullet('Permitindo a geolocalizacao para que o sistema ordene as igrejas pela distancia.')
add_body('Ao selecionar uma igreja no mapa, o codigo dela e preenchido automaticamente no formulario. Caso o membro acesse a pagina atraves de um link de convite (com o parametro ?igreja=CODIGO na URL), a igreja correspondente e pre-selecionada automaticamente no mapa.', indent=True)

add_heading_custom('3.3  Calculo de Distancia', 2)
add_body('Para ordenar as igrejas pela proximidade do usuario, o sistema utiliza a formula de Haversine, que calcula a distancia entre dois pontos na superficie terrestre a partir de suas coordenadas geograficas (latitude e longitude). A distancia e exibida ao lado de cada igreja na lista, em quilometros ou metros conforme a proximidade.', indent=True)

add_heading_custom('3.4  Migracao do Banco de Dados', 2)
add_body('Para suportar a localizacao geografica das igrejas, foi criada uma migracao no banco de dados que adiciona as colunas latitude e longitude a tabela de igrejas. Um indice foi criado para otimizar consultas geograficas. O arquivo de migracao esta em database/migracao-localizacao.sql.', indent=True)

make_table(
    ['Coluna', 'Tipo', 'Descricao'],
    [
        ['latitude', 'DOUBLE PRECISION', 'Latitude da igreja (coordenada geografica)'],
        ['longitude', 'DOUBLE PRECISION', 'Longitude da igreja (coordenada geografica)'],
    ]
)

# ==================== 4 QR CODE ====================
add_heading_custom('4  QR CODE PARA CONVITE DE MEMBROS', 1)
add_body('O painel administrativo da igreja recebeu um recurso de geracao de QR Code que facilita o convite de novos membros. O QR Code codifica um link direto para a pagina de cadastro, ja incluindo o codigo da igreja como parametro na URL.', indent=True)

add_heading_custom('4.1  Funcionamento', 2)
add_body('O QR Code e gerado no proprio navegador utilizando a biblioteca qrcode-generator, carregada via CDN (Content Delivery Network). Nenhuma requisicao ao servidor e necessaria para gerar o codigo. A URL codificada segue o formato:', indent=True)
add_body('https://congregafiel.web.app/autenticacao/criar-conta.html?tipo=membro&igreja=CODIGO', bold=True)
add_body('Quando um membro escaneia o QR Code com o celular, e redirecionado diretamente para a pagina de cadastro com a igreja ja pre-selecionada no mapa, eliminando a necessidade de digitar o codigo manualmente.', indent=True)

add_heading_custom('4.2  Recursos Disponiveis', 2)
add_body('O painel da igreja oferece tres formas de compartilhar o convite:', indent=True)
add_bullet('QR Code visual: exibido diretamente no painel para ser escaneado presencialmente;')
add_bullet('Copiar link: botao que copia o link de convite para a area de transferencia;')
add_bullet('Baixar QR Code: botao que faz o download da imagem PNG do QR Code para impressao ou envio digital.')

# ==================== 5 LINHA DO TEMPO ====================
add_heading_custom('5  LINHA DO TEMPO DE EVENTOS', 1)
add_body('A Linha do Tempo e uma nova pagina que apresenta os eventos da igreja em formato cronologico vertical, oferecendo uma visao completa e organizada da programacao. A funcionalidade foi implementada tanto no painel administrativo quanto no painel dos membros.', indent=True)

add_heading_custom('5.1  Visualizacao Cronologica', 2)
add_body('Os eventos sao organizados por mes, com cada mes formando um grupo visual separado. Uma linha vertical conecta os eventos, com nos (circulos) marcando cada um. Cada evento exibe titulo, data, horario e local, alem de um icone e cor correspondentes ao seu tipo.', indent=True)

make_table(
    ['Tipo de Evento', 'Cor', 'Icone'],
    [
        ['Culto', 'Marrom (#D4A574)', 'Livro aberto'],
        ['Estudo Biblico', 'Azul (#5B8DEF)', 'Livro de estudo'],
        ['Conferencia', 'Roxo (#9B6FD9)', 'Microfone'],
        ['Evento Especial', 'Rosa (#E87C8A)', 'Estrela'],
        ['Outro', 'Bege (#C8956C)', 'Calendario'],
    ]
)

add_heading_custom('5.2  Filtros e Estatisticas', 2)
add_body('A pagina possui botoes de filtro que permitem visualizar apenas eventos de um tipo especifico (cultos, estudos, conferencias ou especiais). No topo, um painel de estatisticas exibe o total de eventos, a quantidade de cultos e o numero de eventos futuros.', indent=True)
add_body('Uma secao de destaque (hero) mostra a contagem regressiva para o proximo evento, exibindo dias, horas e minutos restantes. Essa contagem e atualizada automaticamente.', indent=True)

add_heading_custom('5.3  Animacoes e Responsividade', 2)
add_body('A pagina utiliza a API Intersection Observer do navegador para ativar animacoes de entrada conforme o usuario rola a pagina. Os eventos surgem suavemente ao entrar na area visivel da tela. O layout e responsivo, adaptando-se a dispositivos moveis com reorganizacao dos elementos em coluna unica.', indent=True)

# ==================== 6 CATEGORIZACAO DE EVENTOS ====================
add_heading_custom('6  CATEGORIZACAO DE EVENTOS', 1)
add_body('Os eventos passaram a ser classificados por tipo, permitindo uma organizacao mais clara da programacao da igreja. A categorizacao foi implementada em todas as telas que exibem eventos: pagina de eventos, painel principal e linha do tempo.', indent=True)

add_heading_custom('6.1  Tipos Disponiveis', 2)
add_body('Ao criar ou editar um evento, o administrador seleciona o tipo a partir de uma lista predefinida:', indent=True)

make_table(
    ['Tipo', 'Descricao', 'Exemplo'],
    [
        ['Culto', 'Reunioes regulares de adoracao', 'Culto de Domingo, Culto de Quarta'],
        ['Estudo Biblico', 'Encontros de ensino e estudo', 'Estudo do Livro de Romanos'],
        ['Conferencia', 'Eventos maiores com palestrantes', 'Conferencia de Jovens 2026'],
        ['Evento Especial', 'Celebracoes e ocasioes unicas', 'Aniversario da Igreja'],
        ['Outro', 'Demais atividades', 'Mutirao de Limpeza'],
    ]
)

add_heading_custom('6.2  Migracao do Banco de Dados', 2)
add_body('Para suportar a categorizacao, foi adicionada a coluna tipo a tabela de eventos no banco de dados, com valor padrao "evento". Eventos criados anteriormente recebem o tipo padrao automaticamente. O arquivo de migracao esta em database/migracao-tipo-evento.sql.', indent=True)

# ==================== 7 CARROSSEL DE EVENTOS ====================
add_heading_custom('7  CARROSSEL DE EVENTOS NO PAINEL', 1)
add_body('O painel principal (dashboard) de ambos os perfis - igreja e membro - recebeu um carrossel horizontal de eventos. Esse componente exibe os proximos eventos em cartoes visuais que podem ser navegados por rolagem horizontal ou botoes de seta.', indent=True)

add_heading_custom('7.1  Recursos do Carrossel', 2)
add_body('Cada cartao do carrossel exibe o tipo do evento (com cor correspondente), titulo, data, horario e local. Eventos passados aparecem com opacidade reduzida para diferencia-los visualmente dos eventos futuros. O proximo evento recebe um selo de destaque com animacao pulsante.', indent=True)
add_body('Botoes de filtro acima do carrossel permitem alternar entre todos os eventos ou apenas um tipo especifico. Cada botao possui um indicador colorido correspondente ao tipo do evento.', indent=True)

# ==================== 8 MELHORIAS NA INTERFACE ====================
add_heading_custom('8  MELHORIAS NA INTERFACE E NAVEGACAO', 1)
add_body('Diversas melhorias foram aplicadas a interface do sistema para aprimorar a experiencia do usuario e a consistencia visual.', indent=True)

add_heading_custom('8.1  Menu de Navegacao', 2)
add_body('O menu lateral (sidebar) de ambos os paineis foi atualizado com a inclusao do item "Linha do Tempo", representado por um icone de grafico de linha. O painel de membros tambem recebeu o item "Meu Perfil". Os menus possuem destaque visual no item ativo e um botao de logout posicionado na parte inferior.', indent=True)

add_heading_custom('8.2  Pagina de Perfil do Membro', 2)
add_body('Foi criada uma pagina de perfil para membros que exibe as informacoes pessoais (nome, e-mail), dados da igreja vinculada (nome da igreja, codigo) e a data de cadastro. A pagina permite editar nome e e-mail, alem de oferecer funcionalidade de alteracao de senha com validacao de senha atual.', indent=True)

add_heading_custom('8.3  Filtros nos Pedidos de Oracao', 2)
add_body('A pagina de pedidos de oracao recebeu abas de filtragem por status: Todos, Pendentes, Orados e Respondidos. Os pedidos sao exibidos em grade (grid) e filtrados no lado do cliente conforme a aba selecionada. Um estado vazio e exibido quando nao ha pedidos na categoria selecionada.', indent=True)

add_heading_custom('8.4  Paginas de Comunicados e Pagamentos', 2)
add_body('As paginas de comunicados e pagamentos foram aprimoradas em ambos os paineis. Na visao administrativa, o pastor pode criar, editar e excluir comunicados. Na visao do membro, os comunicados sao exibidos apenas para leitura. A pagina de pagamentos exibe o historico de contribuicoes com resumo do total contribuido.', indent=True)

# ==================== 9 TECNOLOGIAS UTILIZADAS ====================
add_heading_custom('9  TECNOLOGIAS E WEB SERVICES UTILIZADOS', 1)
add_body('Esta sprint envolveu a integracao de tecnologias externas e Web Services gratuitos que agregaram funcionalidades ao sistema sem custo adicional:', indent=True)

make_table(
    ['Tecnologia', 'Tipo', 'Finalidade'],
    [
        ['OpenStreetMap', 'Web Service REST (tiles de mapa)', 'Exibicao de mapa interativo com localizacao das igrejas'],
        ['Leaflet.js', 'Biblioteca JavaScript (CDN)', 'Renderizacao e controle do mapa no navegador'],
        ['Geolocation API', 'API nativa do navegador', 'Obter posicao geografica do usuario'],
        ['qrcode-generator', 'Biblioteca JavaScript (CDN)', 'Geracao de QR Code no lado do cliente'],
        ['Intersection Observer', 'API nativa do navegador', 'Animacoes de entrada na linha do tempo'],
    ]
)

add_body('O OpenStreetMap e um projeto colaborativo que fornece dados cartograficos abertos e gratuitos. Diferente de servicos como Google Maps, nao exige chave de API nem possui limites restritivos de uso para aplicacoes web. A integracao com o Congrega Fiel e feita exclusivamente via protocolo HTTP, caracterizando um consumo de Web Service REST.', indent=True)

# ==================== 10 ENDPOINTS DA API ====================
add_heading_custom('10  NOVOS ENDPOINTS DA API', 1)
add_body('Para suportar as novas funcionalidades, novos endpoints foram adicionados as APIs Express.js e FastAPI:', indent=True)

make_table(
    ['Rota', 'Metodo', 'Descricao'],
    [
        ['/api/igrejas/publicas', 'GET', 'Lista igrejas com coordenadas (sem autenticacao)'],
        ['/api/eventos (campo tipo)', 'POST/PUT', 'Criar e atualizar eventos com categorizacao'],
    ]
)

add_body('O endpoint /api/igrejas/publicas e especialmente importante para a integracao com o mapa, pois permite que a pagina de cadastro carregue a lista de igrejas com suas coordenadas geograficas sem exigir que o usuario esteja autenticado. Os dados retornados incluem apenas informacoes publicas: nome, endereco, codigo, latitude e longitude.', indent=True)

# ==================== 11 ESTRUTURA ATUALIZADA ====================
add_heading_custom('11  ESTRUTURA ATUALIZADA DO PROJETO', 1)
add_body('Com as novas paginas e funcionalidades, a estrutura de diretorios do projeto foi ampliada:', indent=True)

make_table(
    ['Pasta / Arquivo', 'Descricao'],
    [
        ['public/', 'Site completo (HTML, CSS, JavaScript)'],
        ['public/autenticacao/', 'Login, cadastro (com mapa) e recuperacao de senha'],
        ['public/igreja/', 'Painel administrativo do pastor (8 paginas)'],
        ['public/membros/', 'Painel dos membros (8 paginas)'],
        ['public/js/servicos/', 'Servicos compartilhados (sessao, interface, API)'],
        ['api-express/', 'API REST com Express.js + Supabase'],
        ['api-fastapi/', 'API REST com FastAPI + Supabase'],
        ['database/', 'Schema SQL e migracoes do banco de dados'],
    ]
)

add_body('As paginas do painel da igreja agora incluem: Painel, Fieis, Eventos, Contribuicoes, Comunicados, Pedidos de Oracao, Linha do Tempo e configuracoes com QR Code. O painel de membros inclui: Painel, Eventos, Contribuicoes, Comunicados, Pedidos de Oracao, Linha do Tempo e Meu Perfil.', indent=True)

# ==================== 12 CONSIDERACOES FINAIS ====================
add_heading_custom('12  CONSIDERACOES FINAIS', 1)
add_body('A Sprint 3 representou um avanco significativo na experiencia do usuario e na integracao do sistema com servicos web externos. A incorporacao do OpenStreetMap como Web Service gratuito demonstra a capacidade do projeto de consumir APIs externas para agregar funcionalidades sofisticadas sem custos adicionais de infraestrutura.', indent=True)
add_body('O mapa interativo no cadastro de membros transformou um processo antes manual (digitar codigo da igreja) em uma experiencia visual e intuitiva, com suporte a geolocalizacao e busca. O QR Code complementa essa funcionalidade ao permitir que pastores compartilhem convites de forma pratica, seja presencialmente ou por meios digitais.', indent=True)
add_body('A Linha do Tempo trouxe uma nova perspectiva para a visualizacao de eventos, organizando-os cronologicamente com distincao visual por tipo, contagem regressiva e animacoes de rolagem. A categorizacao de eventos por tipo (culto, estudo, conferencia, especial) proporcionou uma organizacao mais clara e permitiu filtragens em todas as telas que exibem eventos.', indent=True)
add_body('As melhorias nos menus de navegacao, a criacao da pagina de perfil do membro e os filtros nos pedidos de oracao consolidaram a usabilidade do sistema. O Congrega Fiel evoluiu de uma plataforma funcional para um sistema com recursos visuais ricos e integracoes modernas, mantendo a filosofia de utilizar exclusivamente tecnologias gratuitas e de codigo aberto.', indent=True)

doc.add_page_break()

# ==================== REFERENCIAS ====================
add_heading_custom('REFERENCIAS', 1)
refs = [
    'LEAFLET. Leaflet - an open-source JavaScript library for mobile-friendly interactive maps. Disponivel em: https://leafletjs.com/. Acesso em: 10 mar. 2026.',
    'OPENSTREETMAP. OpenStreetMap - The Free Wiki World Map. Disponivel em: https://www.openstreetmap.org/. Acesso em: 10 mar. 2026.',
    'MDN WEB DOCS. Geolocation API. Disponivel em: https://developer.mozilla.org/en-US/docs/Web/API/Geolocation_API. Acesso em: 10 mar. 2026.',
    'MDN WEB DOCS. Intersection Observer API. Disponivel em: https://developer.mozilla.org/en-US/docs/Web/API/Intersection_Observer_API. Acesso em: 10 mar. 2026.',
    'QRCODE-GENERATOR. QR Code generator library. Disponivel em: https://github.com/niclas-niclas/qrcode-generator. Acesso em: 10 mar. 2026.',
    'EXPRESS.JS. Express - Node.js web application framework. Disponivel em: https://expressjs.com/. Acesso em: 10 mar. 2026.',
    'FASTAPI. FastAPI - modern, fast web framework for building APIs with Python. Disponivel em: https://fastapi.tiangolo.com/. Acesso em: 10 mar. 2026.',
    'SUPABASE. Supabase - The Open Source Firebase Alternative. Disponivel em: https://supabase.com/docs. Acesso em: 10 mar. 2026.',
    'FIREBASE. Firebase Hosting Documentation. Disponivel em: https://firebase.google.com/docs/hosting. Acesso em: 10 mar. 2026.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

output_path = os.path.join('docs', 'Semana 3.docx')
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
